#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件查看器 - 主程序
支持图形界面和命令行两种模式
"""

import os
import sys

# 尝试导入tkinter，如果失败则使用命令行模式
try:
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox
    from tkinter import font as tkfont

    TKINTER_AVAILABLE = True
except ImportError:
    TKINTER_AVAILABLE = False


class FileViewerApp:
    """文件查看器应用程序类"""

    def __init__(self, root=None):
        """初始化应用程序"""
        self.root = root
        self.is_gui_mode = root is not None

        if self.is_gui_mode:
            self.setup_gui()
        else:
            self.setup_cli()

        # 初始化通用变量
        self.folder_path = ""
        self.file_list = []
        self.selected_files = set()
        self.current_file = None
        self.show_txt = True
        self.show_word = True

    def setup_gui(self):
        """设置图形界面"""
        if not self.root:
            return

        # 设置窗口属性
        self.root.title("文件查看器")
        self.root.geometry("1000x600")
        self.root.minsize(800, 500)

        # 设置中文字体
        try:
            self.default_font = tkfont.nametofont("TkDefaultFont")
            self.default_font.configure(family="SimHei", size=10)
            self.root.option_add("*Font", self.default_font)
        except:
            pass  # 忽略字体设置错误

        # 创建主框架
        self.main_frame = ttk.Frame(self.root)
        self.main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # 创建左侧框架
        self.left_frame = ttk.LabelFrame(self.main_frame, text="文件管理")
        self.left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))

        # 创建右侧框架
        self.right_frame = ttk.LabelFrame(self.main_frame, text="内容预览")
        self.right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 0))

        # 配置左侧框架
        self.setup_left_frame()

        # 配置右侧框架
        self.setup_right_frame()

        # 配置可调整大小的分隔线
        self.paned_window = ttk.PanedWindow(self.main_frame, orient=tk.HORIZONTAL)
        self.paned_window.pack(fill=tk.BOTH, expand=True)
        self.paned_window.add(self.left_frame, weight=1)
        self.paned_window.add(self.right_frame, weight=2)

    def setup_left_frame(self):
        """配置左侧框架"""
        # 创建导入文件夹按钮
        self.import_btn = ttk.Button(self.left_frame, text="导入文件夹", command=self.import_folder)
        self.import_btn.pack(pady=10)

        # 创建文件类型选择框架
        self.type_frame = ttk.Frame(self.left_frame)
        self.type_frame.pack(fill=tk.X, pady=5)

        # 创建TXT文件复选框
        self.txt_var = tk.BooleanVar(value=True)
        self.txt_check = ttk.Checkbutton(self.type_frame, text="识别TXT文件", variable=self.txt_var,
                                         command=self.filter_files)
        self.txt_check.pack(side=tk.LEFT, padx=10)

        # 创建Word文件复选框
        self.word_var = tk.BooleanVar(value=True)
        self.word_check = ttk.Checkbutton(self.type_frame, text="识别Word文件", variable=self.word_var,
                                          command=self.filter_files)
        self.word_check.pack(side=tk.LEFT, padx=10)

        # 创建识别按钮
        self.identify_btn = ttk.Button(self.left_frame, text="识别", command=self.identify_files)
        self.identify_btn.pack(pady=10)

        # 创建文件列表框架
        self.list_frame = ttk.Frame(self.left_frame, bd=2, relief=tk.SUNKEN)
        self.list_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        # 创建文件列表
        self.file_tree = ttk.Treeview(self.list_frame, columns=("checkbox", "filename"), show="headings",
                                      selectmode="none")
        self.file_tree.heading("checkbox", text="")
        self.file_tree.heading("filename", text="文件名")
        self.file_tree.column("checkbox", width=40, anchor=tk.CENTER)
        self.file_tree.column("filename", width=200, anchor=tk.W)

        # 添加滚动条
        self.scrollbar = ttk.Scrollbar(self.list_frame, orient=tk.VERTICAL, command=self.file_tree.yview)
        self.file_tree.configure(yscrollcommand=self.scrollbar.set)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.file_tree.pack(fill=tk.BOTH, expand=True)

        # 绑定双击事件
        self.file_tree.bind("<Double-1>", self.on_file_double_click)

        # 创建删除按钮
        self.delete_btn = ttk.Button(self.left_frame, text="删除选中文件", command=self.delete_selected_files)
        self.delete_btn.pack(pady=10)

    def setup_right_frame(self):
        """配置右侧框架"""
        # 创建文件名预览框架
        self.filename_frame = ttk.LabelFrame(self.right_frame, text="文件名")
        self.filename_frame.pack(fill=tk.X, pady=5)

        # 创建文件名文本框
        self.filename_text = tk.Text(self.filename_frame, height=2, wrap=tk.WORD)
        self.filename_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # 创建复制文件名按钮
        self.copy_filename_btn = ttk.Button(self.filename_frame, text="复制文件名", command=self.copy_filename)
        self.copy_filename_btn.pack(pady=5)

        # 创建文件内容预览框架
        self.content_frame = ttk.LabelFrame(self.right_frame, text="文件内容")
        self.content_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        # 创建文件内容文本框
        self.content_text = tk.Text(self.content_frame, wrap=tk.WORD)
        self.content_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5, side=tk.LEFT)

        # 添加滚动条
        self.content_scrollbar = ttk.Scrollbar(self.content_frame, orient=tk.VERTICAL, command=self.content_text.yview)
        self.content_text.configure(yscrollcommand=self.content_scrollbar.set)
        self.content_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # 创建复制内容按钮
        self.copy_content_btn = ttk.Button(self.right_frame, text="复制内容", command=self.copy_content)
        self.copy_content_btn.pack(pady=5)

    def setup_cli(self):
        """设置命令行界面"""
        self.cli_commands = {
            '1': ('选择文件夹', self.cli_select_folder),
            '2': ('设置文件类型筛选', self.cli_set_file_type_filter),
            '3': ('扫描文件', self.cli_scan_files),
            '4': ('查看文件列表', self.cli_show_file_list),
            '5': ('查看文件内容', self.cli_view_file_content),
            '6': ('删除选中文件', self.cli_delete_selected_files),
            '7': ('退出程序', self.cli_exit)
        }

    def import_folder(self):
        """导入文件夹（GUI版本）"""
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.folder_path = folder_path
            messagebox.showinfo("成功", f"已选择文件夹: {folder_path}")
            # 自动识别文件
            self.identify_files()

    def identify_files(self):
        """识别文件（GUI版本）"""
        if not self.folder_path:
            messagebox.showwarning("警告", "请先导入文件夹")
            return

        # 获取复选框状态
        self.show_txt = self.txt_var.get()
        self.show_word = self.word_var.get()

        # 清空文件列表
        self.file_list = []

        # 扫描文件夹
        try:
            for root, dirs, files in os.walk(self.folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    file_ext = os.path.splitext(file)[1].lower()

                    # 根据选择的文件类型过滤
                    if (self.show_txt and file_ext == '.txt') or \
                            (self.show_word and file_ext == '.docx'):
                        self.file_list.append((file, file_path))

            # 更新文件列表显示
            self.update_file_list()

            if not self.file_list:
                messagebox.showinfo("提示", "未找到符合条件的文件")
            else:
                messagebox.showinfo("成功", f"找到 {len(self.file_list)} 个文件")

        except Exception as e:
            messagebox.showerror("错误", f"扫描文件夹时出错: {str(e)}")

    def filter_files(self):
        """过滤文件（GUI版本）"""
        if self.folder_path:
            self.identify_files()

    def update_file_list(self):
        """更新文件列表显示（GUI版本）"""
        # 清空现有列表
        for item in self.file_tree.get_children():
            self.file_tree.delete(item)

        # 添加文件到列表
        for i, (filename, filepath) in enumerate(self.file_list):
            # 添加到树视图
            self.file_tree.insert("", tk.END, values=("", filename), tags=(str(i),))

            # 创建复选框
            checkbox_var = tk.BooleanVar()
            checkbox = ttk.Checkbutton(self.file_tree, variable=checkbox_var,
                                       command=lambda idx=i: self.toggle_file_selection(idx))

            # 将复选框放置在单元格中
            self.file_tree.after(10, lambda cb=checkbox, idx=i: self.place_checkbox(cb, idx))

    def place_checkbox(self, checkbox, item_index):
        """将复选框放置在树视图单元格中（GUI版本）"""
        # 获取单元格坐标
        x, y, width, height = self.file_tree.bbox(str(item_index), "checkbox")
        if x and y:  # 确保单元格可见
            checkbox.place(in_=self.file_tree, x=x + 5, y=y, anchor=tk.W)

    def toggle_file_selection(self, index):
        """切换文件选择状态"""
        if index in self.selected_files:
            self.selected_files.remove(index)
        else:
            self.selected_files.add(index)

    def on_file_double_click(self, event):
        """双击文件时的操作（GUI版本）"""
        item = self.file_tree.identify_row(event.y)
        if item:
            tags = self.file_tree.item(item, "tags")
            if tags:
                index = int(tags[0])
                self.view_file(index)

    def view_file(self, index):
        """查看文件内容"""
        if 0 <= index < len(self.file_list):
            filename, filepath = self.file_list[index]
            self.current_file = (filename, filepath)

            if self.is_gui_mode:
                # 显示文件名
                self.filename_text.delete(1.0, tk.END)
                self.filename_text.insert(tk.END, filename)

                # 读取并显示文件内容
                try:
                    self.content_text.delete(1.0, tk.END)

                    if filepath.lower().endswith('.txt'):
                        # 尝试多种编码读取文本文件
                        encodings = ['utf-8', 'gbk', 'latin-1']
                        content = None

                        for encoding in encodings:
                            try:
                                with open(filepath, 'r', encoding=encoding) as f:
                                    content = f.read()
                                break
                            except UnicodeDecodeError:
                                continue

                        if content is not None:
                            self.content_text.insert(tk.END, content)
                        else:
                            self.content_text.insert(tk.END, "无法解码文件内容，可能使用了不支持的编码格式。")

                    elif filepath.lower().endswith('.docx'):
                        try:
                            # 尝试导入python-docx库
                            import docx
                            doc = docx.Document(filepath)
                            content = ""
                            for paragraph in doc.paragraphs:
                                content += paragraph.text + '\n'
                            self.content_text.insert(tk.END, content)
                        except ImportError:
                            self.content_text.insert(tk.END, "注意：这是一个Word文档(.docx)。\n\n"
                                                             "由于系统限制，无法直接读取Word文档内容。\n"
                                                             "请安装python-docx库以支持Word文档查看：\n\n"
                                                             "pip install python-docx")
                        except Exception as e:
                            self.content_text.insert(tk.END, f"读取Word文档时出错: {str(e)}")

                except Exception as e:
                    self.content_text.insert(tk.END, f"读取文件时出错: {str(e)}")

            return filename, filepath

    def delete_selected_files(self):
        """删除选中的文件（GUI版本）"""
        if not self.selected_files:
            messagebox.showwarning("警告", "请先选择要删除的文件")
            return

        # 确认删除
        confirm = messagebox.askyesno("确认", f"确定要删除选中的 {len(self.selected_files)} 个文件吗？")
        if not confirm:
            return

        # 删除文件
        self._delete_files()

    def _delete_files(self):
        """执行文件删除操作"""
        deleted_count = 0
        error_count = 0

        # 从后往前删除，避免索引变化问题
        for index in sorted(self.selected_files, reverse=True):
            try:
                filename, filepath = self.file_list[index]
                os.remove(filepath)
                del self.file_list[index]
                deleted_count += 1
            except Exception as e:
                error_count += 1
                if self.is_gui_mode:
                    messagebox.showerror("错误", f"无法删除文件 {filename}: {str(e)}")
                else:
                    print(f"✗ 删除失败 {filename}: {str(e)}")

        # 清空选择
        self.selected_files.clear()

        # 更新文件列表
        if self.is_gui_mode:
            self.update_file_list()
            # 显示结果
            if deleted_count > 0:
                messagebox.showinfo("成功", f"已删除 {deleted_count} 个文件")
        else:
            if deleted_count > 0:
                print(f"\n✓ 已删除 {deleted_count} 个文件")

    def copy_filename(self):
        """复制文件名（GUI版本）"""
        if self.current_file:
            filename, _ = self.current_file
            if self.is_gui_mode:
                self.root.clipboard_clear()
                self.root.clipboard_append(filename)
                messagebox.showinfo("成功", "文件名已复制到剪贴板")
            else:
                # 命令行版本的复制功能
                try:
                    import pyperclip
                    pyperclip.copy(filename)
                    print(f"\n✓ 文件名已复制到剪贴板: {filename}")
                except ImportError:
                    print(f"\n文件名: {filename}")
                    print("提示: 安装pyperclip库可启用复制功能: pip install pyperclip")

    def copy_content(self):
        """复制文件内容（GUI版本）"""
        if self.is_gui_mode:
            content = self.content_text.get(1.0, tk.END)
            if content.strip():
                self.root.clipboard_clear()
                self.root.clipboard_append(content)
                messagebox.showinfo("成功", "文件内容已复制到剪贴板")
            else:
                messagebox.showwarning("警告", "没有可复制的内容")
        else:
            # 命令行版本的复制功能
            if self.current_file:
                filename, filepath = self.current_file
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read()

                    try:
                        import pyperclip
                        pyperclip.copy(content)
                        print(f"\n✓ 文件内容已复制到剪贴板")
                    except ImportError:
                        print(f"\n文件内容:\n{content}")
                        print("\n提示: 安装pyperclip库可启用复制功能: pip install pyperclip")
                except Exception as e:
                    print(f"\n✗ 读取文件内容时出错: {str(e)}")

    def cli_select_folder(self):
        """选择文件夹（命令行版本）"""
        print("\n请输入文件夹路径 (或输入 'q' 返回主菜单):")
        path = input("> ").strip()

        if path.lower() == 'q':
            return

        if os.path.isdir(path):
            self.folder_path = path
            print(f"\n✓ 已选择文件夹: {path}")
        else:
            print(f"\n✗ 错误: '{path}' 不是有效的文件夹路径")

        input("\n按回车键继续...")

    def cli_set_file_type_filter(self):
        """设置文件类型筛选（命令行版本）"""
        print("\n当前设置:")
        print(f"- 显示TXT文件: {'✓' if self.show_txt else '✗'}")
        print(f"- 显示Word文件: {'✓' if self.show_word else '✗'}")
        print()

        print("请选择要修改的选项:")
        print("1. 切换TXT文件显示")
        print("2. 切换Word文件显示")
        print("3. 返回主菜单")

        choice = input("> ").strip()

        if choice == '1':
            self.show_txt = not self.show_txt
            print(f"\n✓ TXT文件显示已{'启用' if self.show_txt else '禁用'}")
        elif choice == '2':
            self.show_word = not self.show_word
            print(f"\n✓ Word文件显示已{'启用' if self.show_word else '禁用'}")
        elif choice == '3':
            return
        else:
            print("\n✗ 无效的选择")

        input("\n按回车键继续...")

    def cli_scan_files(self):
        """扫描文件（命令行版本）"""
        if not self.folder_path:
            print("\n✗ 错误: 请先选择文件夹")
            input("\n按回车键继续...")
            return

        print(f"\n正在扫描文件夹: {self.folder_path}")
        print("扫描中...")

        # 清空文件列表
        self.file_list = []

        try:
            for root, dirs, files in os.walk(self.folder_path):
                for file in files:
                    file_path = os.path.join(root, file)
                    file_ext = os.path.splitext(file)[1].lower()

                    # 根据选择的文件类型过滤
                    if (self.show_txt and file_ext == '.txt') or \
                            (self.show_word and file_ext == '.docx'):
                        self.file_list.append((file, file_path))

            if not self.file_list:
                print("\n✗ 未找到符合条件的文件")
            else:
                print(f"\n✓ 扫描完成，找到 {len(self.file_list)} 个文件")

        except Exception as e:
            print(f"\n✗ 扫描文件夹时出错: {str(e)}")

        input("\n按回车键继续...")

    def cli_show_file_list(self):
        """显示文件列表（命令行版本）"""
        if not self.file_list:
            print("\n✗ 文件列表为空，请先扫描文件")
            input("\n按回车键继续...")
            return

        print(f"\n文件列表 (共 {len(self.file_list)} 个文件):")
        print("-" * 80)
        print(f"{'序号':<5} {'文件名':<40} {'路径':<30}")
        print("-" * 80)

        for i, (filename, filepath) in enumerate(self.file_list, 1):
            status = "✓" if i - 1 in self.selected_files else " "
            print(f"{status} {i:<4} {filename:<40} {os.path.basename(os.path.dirname(filepath)):<30}")

        print("-" * 80)
        print("\n操作选项:")
        print("1. 选择/取消选择文件")
        print("2. 全选")
        print("3. 取消全选")
        print("4. 返回主菜单")

        choice = input("> ").strip()

        if choice == '1':
            self.cli_toggle_file_selection()
        elif choice == '2':
            self.selected_files = set(range(len(self.file_list)))
            print(f"\n✓ 已全选 {len(self.selected_files)} 个文件")
            input("\n按回车键继续...")
        elif choice == '3':
            self.selected_files.clear()
            print("\n✓ 已取消全选")
            input("\n按回车键继续...")
        elif choice == '4':
            return
        else:
            print("\n✗ 无效的选择")
            input("\n按回车键继续...")

    def cli_toggle_file_selection(self):
        """切换文件选择状态（命令行版本）"""
        print("\n请输入要选择/取消选择的文件序号 (多个序号用逗号分隔，如: 1,3,5):")
        input_str = input("> ").strip()

        if not input_str:
            return

        try:
            # 解析输入的序号
            indices = []
            for part in input_str.split(','):
                part = part.strip()
                if '-' in part:  # 处理范围，如 "1-5"
                    start, end = map(int, part.split('-'))
                    indices.extend(range(start, end + 1))
                else:
                    indices.append(int(part))

            # 切换选择状态
            for idx in indices:
                if 1 <= idx <= len(self.file_list):
                    if idx - 1 in self.selected_files:
                        self.selected_files.remove(idx - 1)
                    else:
                        self.selected_files.add(idx - 1)

            print(f"\n✓ 已更新选择状态，当前选中 {len(self.selected_files)} 个文件")

        except ValueError:
            print("\n✗ 无效的输入格式")

        input("\n按回车键继续...")

    def cli_view_file_content(self):
        """查看文件内容（命令行版本）"""
        if not self.file_list:
            print("\n✗ 文件列表为空，请先扫描文件")
            input("\n按回车键继续...")
            return

        print("\n请输入要查看的文件序号 (或输入 'l' 查看文件列表):")
        choice = input("> ").strip()

        if choice.lower() == 'l':
            self.cli_show_file_list()
            return

        try:
            idx = int(choice) - 1

            if 0 <= idx < len(self.file_list):
                filename, filepath = self.view_file(idx)
                print(f"\n文件: {filename}")
                print(f"路径: {filepath}")
                print("-" * 60)

                try:
                    if filepath.lower().endswith('.txt'):
                        # 尝试多种编码读取文本文件
                        encodings = ['utf-8', 'gbk', 'latin-1']
                        content = None

                        for encoding in encodings:
                            try:
                                with open(filepath, 'r', encoding=encoding) as f:
                                    content = f.read()
                                break
                            except UnicodeDecodeError:
                                continue

                        if content is not None:
                            print(content)
                        else:
                            print("无法解码文件内容，可能使用了不支持的编码格式。")

                    elif filepath.lower().endswith('.docx'):
                        try:
                            # 尝试导入python-docx库
                            import docx
                            doc = docx.Document(filepath)
                            print("Word文档内容:")
                            for paragraph in doc.paragraphs:
                                print(paragraph.text)
                        except ImportError:
                            print("注意：这是一个Word文档(.docx)")
                            print("在命令行版本中无法直接查看Word文档内容。")
                            print("建议安装python-docx库以支持Word文档查看：")
                            print("pip install python-docx")
                        except Exception as e:
                            print(f"读取Word文档时出错: {str(e)}")

                except Exception as e:
                    print(f"读取文件时出错: {str(e)}")

                print("-" * 60)
                print("\n操作选项:")
                print("1. 复制文件内容到剪贴板")
                print("2. 返回")

                action = input("> ").strip()

                if action == '1':
                    self.copy_content()
                    input("\n按回车键继续...")

            else:
                print(f"\n✗ 错误: 文件序号 {choice} 不存在")
                input("\n按回车键继续...")

        except ValueError:
            print("\n✗ 错误: 请输入有效的文件序号")
            input("\n按回车键继续...")

    def cli_delete_selected_files(self):
        """删除选中的文件（命令行版本）"""
        if not self.selected_files:
            print("\n✗ 没有选中的文件")
            input("\n按回车键继续...")
            return

        print(f"\n即将删除 {len(self.selected_files)} 个选中的文件:")

        for idx in self.selected_files:
            filename, filepath = self.file_list[idx]
            print(f"  - {filename}")

        print("\n确认删除? (y/n):")
        confirm = input("> ").strip().lower()

        if confirm == 'y':
            self._delete_files()
        else:
            print("\n已取消删除操作")

        input("\n按回车键继续...")

    def cli_exit(self):
        """退出程序（命令行版本）"""
        print("\n感谢使用文件查看器！")
        sys.exit(0)

    def clear_screen(self):
        """清屏（命令行版本）"""
        os.system('cls' if os.name == 'nt' else 'clear')

    def print_header(self):
        """打印标题（命令行版本）"""
        self.clear_screen()
        print("=" * 60)
        print("文件查看器".center(50))
        print("=" * 60)
        print()

    def print_menu(self):
        """打印菜单（命令行版本）"""
        print("请选择操作:")
        for key, (desc, _) in sorted(self.cli_commands.items()):
            print(f"{key}. {desc}")
        print()

    def run_cli(self):
        """运行命令行版本"""
        try:
            while True:
                self.print_header()
                self.print_menu()

                choice = input("> ").strip()

                if choice in self.cli_commands:
                    _, func = self.cli_commands[choice]
                    func()
                else:
                    print("\n✗ 无效的选择，请重新输入")
                    input("\n按回车键继续...")

        except KeyboardInterrupt:
            print("\n\n程序被用户中断")
        except Exception as e:
            print(f"\n✗ 程序发生错误: {str(e)}")
            input("\n按回车键退出...")

    def run_gui(self):
        """运行图形界面版本"""
        if not self.is_gui_mode:
            print("错误: 图形界面版本需要提供Tkinter根窗口")
            return

        try:
            self.root.mainloop()
        except KeyboardInterrupt:
            print("\n程序被用户中断")
        except Exception as e:
            print(f"\n程序发生错误: {str(e)}")

    def run(self):
        """运行程序"""
        if self.is_gui_mode:
            self.run_gui()
        else:
            self.run_cli()


def main():
    """主函数"""
    # 检查是否支持图形界面
    if TKINTER_AVAILABLE:
        try:
            root = tk.Tk()
            print("启动图形界面版本...")
            app = FileViewerApp(root)
            app.run()
        except Exception as e:
            print(f"图形界面启动失败: {str(e)}")
            print("切换到命令行版本...")
            app = FileViewerApp()
            app.run()
    else:
        print("未检测到图形界面支持，启动命令行版本...")
        app = FileViewerApp()
        app.run()


if __name__ == "__main__":
    main()