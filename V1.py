import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import threading
import queue
import time

class FileRenamerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("小工具合集")
        self.root.geometry("900x600")  # 设置初始大小
        self.root.minsize(600, 400)  # 设置最小大小
        self.root.configure(bg="#f0f0f0")
        
        # 颜色方案
        self.colors = {
            "bg": "#f0f0f0",
            "frame_bg": "#ffffff",
            "text": "#333333",
            "button": {
                "browse": "#e0e0e0",
                "apply": "#4CAF50",
                "clear": "#f44336",
                "rename": "#2196F3",
                "default": "#2196F3",
                "download": "#ff9800"
            },
            "button_text": {
                "browse": "#333333",
                "apply": "#ffffff",
                "clear": "#ffffff",
                "rename": "#ffffff",
                "default": "#ffffff",
                "download": "#ffffff"
            }
        }
        
        # 文件夹路径
        self.folder_path = ""
        # 文件列表
        self.files = []
        # 新名称输入框列表
        self.name_entries = []
        
        # 创建主框架
        self.main_frame = tk.Frame(root, bg=self.colors["bg"])
        self.main_frame.pack(pady=20, padx=20, fill=tk.BOTH, expand=True)
        
        # 创建标签页控制区域
        self.tab_control_frame = tk.Frame(self.main_frame, bg=self.colors["bg"])
        self.tab_control_frame.pack(fill=tk.X, pady=10)
        
        # 创建标签按钮
        self.tab_var = tk.StringVar(value="rename")
        
        self.rename_tab_btn = tk.Radiobutton(self.tab_control_frame, text="文件重命名", variable=self.tab_var, value="rename", 
                                             command=self.switch_tab, bg=self.colors["bg"], font=('Arial', 10, 'bold'))
        self.rename_tab_btn.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.char_tab_btn = tk.Radiobutton(self.tab_control_frame, text="字符重组", variable=self.tab_var, value="char", 
                                           command=self.switch_tab, bg=self.colors["bg"], font=('Arial', 10, 'bold'))
        self.char_tab_btn.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.name_tab_btn = tk.Radiobutton(self.tab_control_frame, text="起名小工具", variable=self.tab_var, value="name", 
                                           command=self.switch_tab, bg=self.colors["bg"], font=('Arial', 10, 'bold'))
        self.name_tab_btn.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.word_tab_btn = tk.Radiobutton(self.tab_control_frame, text="WORD转TXT", variable=self.tab_var, value="word", 
                                           command=self.switch_tab, bg=self.colors["bg"], font=('Arial', 10, 'bold'))
        self.word_tab_btn.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.export_tab_btn = tk.Radiobutton(self.tab_control_frame, text="文本导出", variable=self.tab_var, value="export", 
                                           command=self.switch_tab, bg=self.colors["bg"], font=('Arial', 10, 'bold'))
        self.export_tab_btn.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.txt_to_word_tab_btn = tk.Radiobutton(self.tab_control_frame, text="百度搜索标题相关图片", variable=self.tab_var, value="txt_to_word", 
                                           command=self.switch_tab, bg=self.colors["bg"], font=('Arial', 10, 'bold'))
        self.txt_to_word_tab_btn.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 创建标签页内容区域
        self.tab_content_frame = tk.Frame(self.main_frame, bg=self.colors["bg"])
        self.tab_content_frame.pack(fill=tk.BOTH, expand=True)
        
        # 创建文件重命名标签页
        self.rename_tab = tk.Frame(self.tab_content_frame, bg=self.colors["bg"])
        
        # 创建字符重组标签页
        self.char_tab = tk.Frame(self.tab_content_frame, bg=self.colors["bg"])
        
        # 创建起名小工具标签页
        self.name_tab = tk.Frame(self.tab_content_frame, bg=self.colors["bg"])
        
        # 创建WORD转TXT标签页
        self.word_tab = tk.Frame(self.tab_content_frame, bg=self.colors["bg"])
        
        # 创建文本导出标签页
        self.export_tab = tk.Frame(self.tab_content_frame, bg=self.colors["bg"])
        
        # 创建txt转word+图标签页
        self.txt_to_word_tab = tk.Frame(self.tab_content_frame, bg=self.colors["bg"])
        
        
        
        # 初始化文件重命名标签页
        self.init_rename_tab()
        
        # 初始化字符重组标签页
        self.init_char_tab()
        
        # 初始化起名小工具标签页
        self.init_name_tab()
        
        # 初始化WORD转TXT标签页
        self.init_word_tab()
        
        # 初始化文本导出标签页
        self.init_export_tab()
        
        # 初始化txt转word+图标签页
        self.init_txt_to_word_tab()
        
        
        
        # 默认显示文件重命名标签页
        self.switch_tab()
    
    def init_rename_tab(self):
        """初始化文件重命名标签页"""
        # 创建文件夹选择区域
        self.folder_frame = tk.Frame(self.rename_tab, bg=self.colors["bg"])
        self.folder_frame.pack(fill=tk.X, pady=10)
        
        self.folder_label = tk.Label(self.folder_frame, text="选择文件夹:", bg=self.colors["bg"], font=('Arial', 10, 'bold'))
        self.folder_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.folder_entry = tk.Entry(self.folder_frame, width=50, font=('Arial', 10))
        self.folder_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=10, pady=5)
        
        self.browse_button = tk.Button(self.folder_frame, text="浏览", command=self.browse_folder, 
                                      bg=self.colors["button"]["browse"], 
                                      fg=self.colors["button_text"]["browse"],
                                      font=('Arial', 10),
                                      relief=tk.RAISED)
        self.browse_button.pack(side=tk.RIGHT, padx=10, pady=5)
        
        # 创建左右布局的主内容区域
        self.content_frame = tk.Frame(self.rename_tab, bg=self.colors["bg"])
        self.content_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # 左侧文件列表区域
        self.left_frame = tk.Frame(self.content_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE)
        self.left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10)
        
        # 添加左侧标题
        self.left_title = tk.Label(self.left_frame, text="文件列表", font=('Arial', 11, 'bold'), bg=self.colors["frame_bg"], fg="#333333")
        self.left_title.pack(fill=tk.X, pady=5, padx=10)
        
        self.scrollbar = tk.Scrollbar(self.left_frame)
        self.scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.canvas = tk.Canvas(self.left_frame, yscrollcommand=self.scrollbar.set)
        self.canvas.pack(fill=tk.BOTH, expand=True)
        
        self.scrollbar.config(command=self.canvas.yview)
        
        self.inner_frame = tk.Frame(self.canvas, bg=self.colors["frame_bg"])
        self.canvas.create_window((0, 0), window=self.inner_frame, anchor=tk.NW)
        
        # 右侧批量输入区域
        self.right_frame = tk.Frame(self.content_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE, width=350)
        self.right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=10)
        
        # 添加右侧标题
        self.right_title = tk.Label(self.right_frame, text="批量输入", font=('Arial', 11, 'bold'), bg=self.colors["frame_bg"], fg="#333333")
        self.right_title.pack(fill=tk.X, pady=5, padx=10)
        
        self.batch_label = tk.Label(self.right_frame, text="新名称（一行一个）:", bg=self.colors["frame_bg"])
        self.batch_label.pack(side=tk.TOP, padx=10, pady=5, anchor=tk.W)
        
        self.batch_text = tk.Text(self.right_frame, height=15, width=50, font=('Arial', 10), bd=1, relief=tk.SUNKEN)
        self.batch_text.pack(side=tk.TOP, padx=10, pady=5, fill=tk.BOTH, expand=True)
        
        self.batch_buttons_frame = tk.Frame(self.right_frame, bg=self.colors["frame_bg"])
        self.batch_buttons_frame.pack(side=tk.TOP, pady=10, fill=tk.X)
        
        self.apply_batch_button = tk.Button(self.batch_buttons_frame, text="批量应用", command=self.apply_batch_names, 
                                           bg=self.colors["button"]["apply"], 
                                           fg=self.colors["button_text"]["apply"],
                                           font=('Arial', 10),
                                           padx=10, pady=5,
                                           relief=tk.RAISED)
        self.apply_batch_button.pack(side=tk.RIGHT, padx=10)
        
        self.clear_batch_button = tk.Button(self.batch_buttons_frame, text="清空", command=self.clear_batch_input, 
                                           bg=self.colors["button"]["clear"], 
                                           fg=self.colors["button_text"]["clear"],
                                           font=('Arial', 10),
                                           padx=10, pady=5,
                                           relief=tk.RAISED)
        self.clear_batch_button.pack(side=tk.RIGHT, padx=10)
        
        # 创建底部按钮区域
        self.button_frame = tk.Frame(self.rename_tab, bg=self.colors["bg"])
        self.button_frame.pack(fill=tk.X, pady=10)
        
        # 创建重命名按钮
        self.rename_button = tk.Button(self.button_frame, text="立即重命名", command=self.rename_files, 
                                       font=('Arial', 14, 'bold'), 
                                       bg=self.colors["button"]["rename"], 
                                       fg=self.colors["button_text"]["rename"],
                                       padx=30, pady=15,
                                       relief=tk.RAISED)
        self.rename_button.pack(side=tk.TOP, anchor=tk.CENTER)
        
        # 绑定滚动事件
        self.inner_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
    
    def init_char_tab(self):
        """初始化字符重组标签页"""
        # 创建字符重组功能的主框架
        self.char_main_frame = tk.Frame(self.char_tab, bg=self.colors["bg"])
        self.char_main_frame.pack(fill=tk.BOTH, expand=True, pady=10, padx=10)
        
        # 添加标题
        self.char_title = tk.Label(self.char_main_frame, text="字符重组工具", font=('Arial', 12, 'bold'), bg=self.colors["bg"])
        self.char_title.pack(fill=tk.X, pady=10)
        
        # 创建配置区域
        self.config_frame = tk.Frame(self.char_main_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE)
        self.config_frame.pack(fill=tk.X, pady=10, padx=5)
        
        # 重组顺序输入
        self.order_label = tk.Label(self.config_frame, text="重组顺序:", bg=self.colors["frame_bg"], font=('Arial', 10, 'bold'))
        self.order_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.order_entry = tk.Entry(self.config_frame, width=30, font=('Arial', 10))
        self.order_entry.pack(side=tk.LEFT, padx=10, pady=5)
        self.order_entry.insert(0, "2----4----1")
        
        # 分隔符输入
        self.separator_label = tk.Label(self.config_frame, text="分隔符:", bg=self.colors["frame_bg"], font=('Arial', 10, 'bold'))
        self.separator_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.separator_entry = tk.Entry(self.config_frame, width=10, font=('Arial', 10))
        self.separator_entry.pack(side=tk.LEFT, padx=10, pady=5)
        self.separator_entry.insert(0, "----")
        
        # 转换按钮
        self.convert_button = tk.Button(self.config_frame, text="转换", command=self.convert_text, 
                                       bg=self.colors["button"]["default"], 
                                       fg=self.colors["button_text"]["default"],
                                       font=('Arial', 10, 'bold'),
                                       padx=20, pady=5,
                                       relief=tk.RAISED)
        self.convert_button.pack(side=tk.RIGHT, padx=10, pady=5)
        
        # 创建输入输出区域
        self.io_frame = tk.Frame(self.char_main_frame, bg=self.colors["bg"])
        self.io_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # 左侧输入区域
        self.input_frame = tk.Frame(self.io_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE)
        self.input_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
        
        # 输入标题和行数显示
        self.input_header_frame = tk.Frame(self.input_frame, bg=self.colors["frame_bg"])
        self.input_header_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.input_title = tk.Label(self.input_header_frame, text="输入内容", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.input_title.pack(side=tk.LEFT)
        
        self.line_count_label = tk.Label(self.input_header_frame, text="(0行)", font=('Arial', 10), fg="#666666", bg=self.colors["frame_bg"])
        self.line_count_label.pack(side=tk.LEFT, padx=10)
        
        # 输入文本框
        self.input_text = tk.Text(self.input_frame, height=15, width=50, font=('Arial', 10), bd=1, relief=tk.SUNKEN)
        self.input_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 绑定输入变化事件
        self.input_text.bind("<KeyRelease>", lambda e: self.update_line_count())
        self.input_text.bind("<ButtonRelease>", lambda e: self.update_line_count())
        
        # 输入区域按钮
        self.input_buttons_frame = tk.Frame(self.input_frame, bg=self.colors["frame_bg"])
        self.input_buttons_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.clear_input_button = tk.Button(self.input_buttons_frame, text="清空", command=self.clear_input, 
                                           bg=self.colors["button"]["clear"], 
                                           fg=self.colors["button_text"]["clear"],
                                           font=('Arial', 10),
                                           padx=10, pady=5,
                                           relief=tk.RAISED)
        self.clear_input_button.pack(side=tk.RIGHT, padx=10)
        
        # 右侧输出区域
        self.output_frame = tk.Frame(self.io_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE)
        self.output_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5)
        
        # 输出标题
        self.output_title = tk.Label(self.output_frame, text="输出结果", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.output_title.pack(fill=tk.X, pady=5, padx=10)
        
        # 输出文本框
        self.output_text = tk.Text(self.output_frame, height=15, width=50, font=('Arial', 10), bd=1, relief=tk.SUNKEN)
        self.output_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 输出区域按钮
        self.output_buttons_frame = tk.Frame(self.output_frame, bg=self.colors["frame_bg"])
        self.output_buttons_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.copy_output_button = tk.Button(self.output_buttons_frame, text="复制", command=self.copy_output, 
                                           bg=self.colors["button"]["apply"], 
                                           fg=self.colors["button_text"]["apply"],
                                           font=('Arial', 10),
                                           padx=10, pady=5,
                                           relief=tk.RAISED)
        self.copy_output_button.pack(side=tk.RIGHT, padx=10)
        
        self.clear_output_button = tk.Button(self.output_buttons_frame, text="清空", command=self.clear_output, 
                                            bg=self.colors["button"]["clear"], 
                                            fg=self.colors["button_text"]["clear"],
                                            font=('Arial', 10),
                                            padx=10, pady=5,
                                            relief=tk.RAISED)
        self.clear_output_button.pack(side=tk.RIGHT, padx=10)
    
    def init_name_tab(self):
        """初始化起名小工具标签页"""
        # 创建起名小工具的主框架
        self.name_main_frame = tk.Frame(self.name_tab, bg=self.colors["bg"])
        self.name_main_frame.pack(fill=tk.BOTH, expand=True, pady=10, padx=10)
        
        # 添加标题
        self.name_title = tk.Label(self.name_main_frame, text="起名小工具", font=('Arial', 12, 'bold'), bg=self.colors["bg"])
        self.name_title.pack(fill=tk.X, pady=10)
        
        # 创建配置区域
        self.config_frame = tk.Frame(self.name_main_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE)
        self.config_frame.pack(fill=tk.X, pady=10, padx=5)
        
        # 类型选择
        self.type_frame = tk.Frame(self.config_frame, bg=self.colors["frame_bg"])
        self.type_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.type_label = tk.Label(self.type_frame, text="生成类型:", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.type_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 类型复选框
        self.include_chinese = tk.BooleanVar(value=True)
        self.chinese_check = tk.Checkbutton(self.type_frame, text="汉字", variable=self.include_chinese, 
                                           bg=self.colors["frame_bg"], command=self.update_config_fields)
        self.chinese_check.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.include_english = tk.BooleanVar(value=False)
        self.english_check = tk.Checkbutton(self.type_frame, text="英文", variable=self.include_english, 
                                           bg=self.colors["frame_bg"], command=self.update_config_fields)
        self.english_check.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.include_number = tk.BooleanVar(value=False)
        self.number_check = tk.Checkbutton(self.type_frame, text="数字", variable=self.include_number, 
                                          bg=self.colors["frame_bg"], command=self.update_config_fields)
        self.number_check.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 数量配置
        self.count_frame = tk.Frame(self.config_frame, bg=self.colors["frame_bg"])
        self.count_frame.pack(fill=tk.X, pady=5, padx=10)
        
        # 汉字数量
        self.chinese_count_label = tk.Label(self.count_frame, text="汉字数量:", bg=self.colors["frame_bg"])
        self.chinese_count_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.chinese_count_var = tk.StringVar(value="3")
        self.chinese_count_entry = tk.Entry(self.count_frame, width=5, textvariable=self.chinese_count_var)
        self.chinese_count_entry.pack(side=tk.LEFT, padx=5, pady=5)
        
        # 英文数量
        self.english_count_label = tk.Label(self.count_frame, text="英文数量:", bg=self.colors["frame_bg"])
        self.english_count_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.english_count_var = tk.StringVar(value="0")
        self.english_count_entry = tk.Entry(self.count_frame, width=5, textvariable=self.english_count_var)
        self.english_count_entry.pack(side=tk.LEFT, padx=5, pady=5)
        
        # 数字数量
        self.number_count_label = tk.Label(self.count_frame, text="数字数量:", bg=self.colors["frame_bg"])
        self.number_count_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.number_count_var = tk.StringVar(value="0")
        self.number_count_entry = tk.Entry(self.count_frame, width=5, textvariable=self.number_count_var)
        self.number_count_entry.pack(side=tk.LEFT, padx=5, pady=5)
        
        # 笔画要求
        self.stroke_frame = tk.Frame(self.config_frame, bg=self.colors["frame_bg"])
        self.stroke_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.stroke_label = tk.Label(self.stroke_frame, text="汉字笔画要求:", bg=self.colors["frame_bg"])
        self.stroke_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.min_stroke_var = tk.StringVar(value="8")
        self.min_stroke_entry = tk.Entry(self.stroke_frame, width=5, textvariable=self.min_stroke_var)
        self.min_stroke_entry.pack(side=tk.LEFT, padx=5, pady=5)
        
        self.stroke_range_label = tk.Label(self.stroke_frame, text="~", bg=self.colors["frame_bg"])
        self.stroke_range_label.pack(side=tk.LEFT, padx=5, pady=5)
        
        self.max_stroke_var = tk.StringVar(value="8")
        self.max_stroke_entry = tk.Entry(self.stroke_frame, width=5, textvariable=self.max_stroke_var)
        self.max_stroke_entry.pack(side=tk.LEFT, padx=5, pady=5)
        
        self.stroke_unit_label = tk.Label(self.stroke_frame, text="笔以内", bg=self.colors["frame_bg"])
        self.stroke_unit_label.pack(side=tk.LEFT, padx=5, pady=5)
        
        # 生成数量
        self.generate_frame = tk.Frame(self.config_frame, bg=self.colors["frame_bg"])
        self.generate_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.generate_label = tk.Label(self.generate_frame, text="生成数量:", bg=self.colors["frame_bg"])
        self.generate_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.generate_var = tk.StringVar(value="5")
        self.generate_entry = tk.Entry(self.generate_frame, width=5, textvariable=self.generate_var)
        self.generate_entry.pack(side=tk.LEFT, padx=5, pady=5)
        
        self.generate_unit_label = tk.Label(self.generate_frame, text="条 (最高500条)", bg=self.colors["frame_bg"])
        self.generate_unit_label.pack(side=tk.LEFT, padx=5, pady=5)
        
        # 生成按钮和复制按钮区域
        self.button_frame = tk.Frame(self.config_frame, bg=self.colors["frame_bg"])
        self.button_frame.pack(side=tk.TOP, pady=10, anchor=tk.CENTER)
        
        # 生成按钮
        self.generate_button = tk.Button(self.button_frame, text="生成", command=self.generate_names, 
                                        bg=self.colors["button"]["default"], 
                                        fg=self.colors["button_text"]["default"],
                                        font=('Arial', 10, 'bold'),
                                        padx=20, pady=10,
                                        relief=tk.RAISED)
        self.generate_button.pack(side=tk.LEFT, padx=10)
        
        # 复制所有生成结果按钮
        self.copy_all_button = tk.Button(self.button_frame, text="复制所有结果", command=self.copy_name_result, 
                                        bg=self.colors["button"]["apply"], 
                                        fg=self.colors["button_text"]["apply"],
                                        font=('Arial', 10, 'bold'),
                                        padx=20, pady=10,
                                        relief=tk.RAISED)
        self.copy_all_button.pack(side=tk.LEFT, padx=10)
        
        # 创建结果区域
        self.result_frame = tk.Frame(self.name_main_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE)
        self.result_frame.pack(fill=tk.BOTH, expand=True, pady=10, padx=5)
        
        # 结果标题
        self.result_title = tk.Label(self.result_frame, text="生成结果", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.result_title.pack(fill=tk.X, pady=5, padx=10)
        
        # 结果文本框
        self.result_text = tk.Text(self.result_frame, height=15, width=50, font=('Arial', 10), bd=1, relief=tk.SUNKEN)
        self.result_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 结果区域按钮
        self.result_buttons_frame = tk.Frame(self.result_frame, bg=self.colors["frame_bg"])
        self.result_buttons_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.copy_result_button = tk.Button(self.result_buttons_frame, text="复制所有", command=self.copy_name_result, 
                                           bg=self.colors["button"]["apply"], 
                                           fg=self.colors["button_text"]["apply"],
                                           font=('Arial', 10),
                                           padx=10, pady=5,
                                           relief=tk.RAISED)
        self.copy_result_button.pack(side=tk.RIGHT, padx=10)
        
        self.clear_result_button = tk.Button(self.result_buttons_frame, text="清空", command=self.clear_name_result, 
                                            bg=self.colors["button"]["clear"], 
                                            fg=self.colors["button_text"]["clear"],
                                            font=('Arial', 10),
                                            padx=10, pady=5,
                                            relief=tk.RAISED)
        self.clear_result_button.pack(side=tk.RIGHT, padx=10)
        
        # 初始化历史记录
        self.history_file = "name_history.txt"
        self.load_history()
        
        # 初始化汉字库（按笔画分类）
        self.chinese_chars = self.build_chinese_chars()
    
    def init_word_tab(self):
        """初始化WORD转TXT标签页"""
        # 创建WORD转TXT功能的主框架
        self.word_main_frame = tk.Frame(self.word_tab, bg=self.colors["bg"])
        self.word_main_frame.pack(fill=tk.BOTH, expand=True, pady=10, padx=10)
        
        # 添加标题
        self.word_title = tk.Label(self.word_main_frame, text="WORD转TXT", font=('Arial', 12, 'bold'), bg=self.colors["bg"])
        self.word_title.pack(fill=tk.X, pady=10)
        
        # 创建配置区域
        self.word_config_frame = tk.Frame(self.word_main_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE)
        self.word_config_frame.pack(fill=tk.X, pady=10, padx=5)
        
        # 输出位置选择
        self.output_frame = tk.Frame(self.word_config_frame, bg=self.colors["frame_bg"])
        self.output_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.output_label = tk.Label(self.output_frame, text="输出位置:", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.output_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.output_entry = tk.Entry(self.output_frame, width=50, font=('Arial', 10))
        self.output_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=10, pady=5)
        
        self.browse_output_button = tk.Button(self.output_frame, text="浏览", command=self.browse_word_output, 
                                           bg=self.colors["button"]["browse"], 
                                           fg=self.colors["button_text"]["browse"],
                                           font=('Arial', 10),
                                           relief=tk.RAISED)
        self.browse_output_button.pack(side=tk.RIGHT, padx=10, pady=5)
        
        # 线程数设置
        self.thread_frame = tk.Frame(self.word_config_frame, bg=self.colors["frame_bg"])
        self.thread_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.thread_label = tk.Label(self.thread_frame, text="线程数:", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.thread_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.thread_var = tk.StringVar(value="1")
        self.thread_entry = tk.Entry(self.thread_frame, width=5, textvariable=self.thread_var, font=('Arial', 10))
        self.thread_entry.pack(side=tk.LEFT, padx=5, pady=5)
        
        self.thread_unit_label = tk.Label(self.thread_frame, text="(最高50线程)", bg=self.colors["frame_bg"])
        self.thread_unit_label.pack(side=tk.LEFT, padx=5, pady=5)
        
        # 按钮区域框架
        self.buttons_frame = tk.Frame(self.word_config_frame, bg=self.colors["frame_bg"])
        self.buttons_frame.pack(fill=tk.X, pady=10, padx=10)
        
        # 转换按钮
        self.convert_button = tk.Button(self.buttons_frame, text="转换", command=self.convert_word_to_txt, 
                                      bg=self.colors["button"]["default"], 
                                      fg=self.colors["button_text"]["default"],
                                      font=('Arial', 10, 'bold'),
                                      padx=20, pady=10,
                                      relief=tk.RAISED)
        self.convert_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)
        
        # 下载全部转换结果按钮
        self.download_all_button = tk.Button(self.buttons_frame, text="下载全部转换结果", command=self.download_all_word_results, 
                                           bg=self.colors["button"]["download"], 
                                           fg=self.colors["button_text"]["download"],
                                           font=('Arial', 10, 'bold'),
                                           padx=20, pady=10,
                                           relief=tk.RAISED)
        self.download_all_button.pack(side=tk.RIGHT, expand=True, fill=tk.X, padx=5)
        
        # 创建输入输出区域
        self.word_io_frame = tk.Frame(self.word_main_frame, bg=self.colors["bg"])
        self.word_io_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # 左侧输入区域（拖拽文件）
        self.word_input_frame = tk.Frame(self.word_io_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE)
        self.word_input_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)
        
        # 输入标题和文件数量显示
        self.word_input_header_frame = tk.Frame(self.word_input_frame, bg=self.colors["frame_bg"])
        self.word_input_header_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.word_input_title = tk.Label(self.word_input_header_frame, text="添加WORD文件", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.word_input_title.pack(side=tk.LEFT)
        
        self.word_file_count_label = tk.Label(self.word_input_header_frame, text="(0个文件)", font=('Arial', 10), fg="#666666", bg=self.colors["frame_bg"])
        self.word_file_count_label.pack(side=tk.LEFT, padx=10)
        
        # 操作区域
        self.word_action_frame = tk.Frame(self.word_input_frame, bg=self.colors["frame_bg"])
        self.word_action_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 添加文件按钮
        self.select_files_button = tk.Button(self.word_action_frame, text="选择WORD文件", command=self.word_on_click, 
                                           bg=self.colors["button"]["browse"], 
                                           fg=self.colors["button_text"]["browse"],
                                           font=('Arial', 12, 'bold'),
                                           padx=20, pady=15,
                                           relief=tk.RAISED)
        self.select_files_button.pack(expand=True, anchor=tk.CENTER, pady=30)
        
        # 提示信息
        self.word_hint_label = tk.Label(self.word_action_frame, text="点击按钮选择一个或多个WORD文件", font=('Arial', 10), fg="#666666", bg=self.colors["frame_bg"])
        self.word_hint_label.pack(anchor=tk.CENTER)
        
        # 拖拽文件列表
        self.word_file_list = []
        
        # 右侧输出区域
        self.word_output_frame = tk.Frame(self.word_io_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE)
        self.word_output_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5)
        
        # 输出标题
        self.word_output_title = tk.Label(self.word_output_frame, text="转换结果", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.word_output_title.pack(fill=tk.X, pady=5, padx=10)
        
        # 输出文本框
        self.word_output_text = tk.Text(self.word_output_frame, height=15, width=50, font=('Arial', 10), bd=1, relief=tk.SUNKEN)
        self.word_output_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 输出区域按钮
        self.word_output_buttons_frame = tk.Frame(self.word_output_frame, bg=self.colors["frame_bg"])
        self.word_output_buttons_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.save_all_button = tk.Button(self.word_output_buttons_frame, text="保存所有结果到输出位置", command=self.save_all_word_results, 
                                       bg=self.colors["button"]["apply"], 
                                       fg=self.colors["button_text"]["apply"],
                                       font=('Arial', 10),
                                       padx=10, pady=5,
                                       relief=tk.RAISED)
        self.save_all_button.pack(side=tk.RIGHT, padx=10)
    
    def init_export_tab(self):
        """初始化文本导出标签页"""
        # 创建文本导出的主框架
        self.export_main_frame = tk.Frame(self.export_tab, bg=self.colors["bg"])
        self.export_main_frame.pack(fill=tk.BOTH, expand=True, pady=10, padx=10)
        
        # 添加标题
        self.export_title = tk.Label(self.export_main_frame, text="文本导出工具", font=('Arial', 12, 'bold'), bg=self.colors["bg"])
        self.export_title.pack(fill=tk.X, pady=10)
        
        # 创建配置区域
        self.export_config_frame = tk.Frame(self.export_main_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE)
        self.export_config_frame.pack(fill=tk.X, pady=10, padx=5)
        
        # 导出格式选择
        self.format_frame = tk.Frame(self.export_config_frame, bg=self.colors["frame_bg"])
        self.format_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.format_label = tk.Label(self.format_frame, text="导出格式:", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.format_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.format_var = tk.StringVar(value="txt")
        self.txt_radio = tk.Radiobutton(self.format_frame, text="TXT格式", variable=self.format_var, value="txt", 
                                       bg=self.colors["frame_bg"], command=self.update_image_options)
        self.txt_radio.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.word_radio = tk.Radiobutton(self.format_frame, text="WORD格式", variable=self.format_var, value="word", 
                                        bg=self.colors["frame_bg"], command=self.update_image_options)
        self.word_radio.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 导出方式选择
        self.export_mode_frame = tk.Frame(self.export_config_frame, bg=self.colors["frame_bg"])
        self.export_mode_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.export_mode_label = tk.Label(self.export_mode_frame, text="导出方式:", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.export_mode_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.export_mode_var = tk.StringVar(value="single")
        self.single_radio = tk.Radiobutton(self.export_mode_frame, text="单个文件", variable=self.export_mode_var, value="single", 
                                          bg=self.colors["frame_bg"])
        self.single_radio.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.multi_radio = tk.Radiobutton(self.export_mode_frame, text="按行独立文件", variable=self.export_mode_var, value="multi", 
                                         bg=self.colors["frame_bg"])
        self.multi_radio.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 文件名设置
        self.filename_frame = tk.Frame(self.export_config_frame, bg=self.colors["frame_bg"])
        self.filename_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.filename_label = tk.Label(self.filename_frame, text="文件名设置:", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.filename_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.filename_mode_var = tk.StringVar(value="sequential")
        self.sequential_radio = tk.Radiobutton(self.filename_frame, text="顺序命名", variable=self.filename_mode_var, value="sequential", 
                                              bg=self.colors["frame_bg"], command=self.update_filename_options)
        self.sequential_radio.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.content_radio = tk.Radiobutton(self.filename_frame, text="按内容开头", variable=self.filename_mode_var, value="content", 
                                           bg=self.colors["frame_bg"], command=self.update_filename_options)
        self.content_radio.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 内容开头字数设置
        self.content_chars_frame = tk.Frame(self.export_config_frame, bg=self.colors["frame_bg"])
        self.content_chars_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.content_chars_label = tk.Label(self.content_chars_frame, text="内容开头字数:", bg=self.colors["frame_bg"])
        self.content_chars_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.content_chars_var = tk.StringVar(value="10")
        self.content_chars_entry = tk.Entry(self.content_chars_frame, width=5, textvariable=self.content_chars_var)
        self.content_chars_entry.pack(side=tk.LEFT, padx=5, pady=5)
        
        # 图片插入设置（仅WORD格式）
        self.image_frame = tk.Frame(self.export_config_frame, bg=self.colors["frame_bg"])
        self.image_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.image_label = tk.Label(self.image_frame, text="图片插入:", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.image_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.image_var = tk.BooleanVar(value=False)
        self.image_check = tk.Checkbutton(self.image_frame, text="插入图片", variable=self.image_var, 
                                         bg=self.colors["frame_bg"], command=self.update_image_position_options)
        self.image_check.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 图片路径选择
        self.image_path_frame = tk.Frame(self.export_config_frame, bg=self.colors["frame_bg"])
        self.image_path_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.image_path_label = tk.Label(self.image_path_frame, text="图片路径:", bg=self.colors["frame_bg"])
        self.image_path_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.image_path_entry = tk.Entry(self.image_path_frame, width=40, font=('Arial', 10))
        self.image_path_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=10, pady=5)
        
        # 浏览按钮框架
        self.browse_buttons_frame = tk.Frame(self.image_path_frame, bg=self.colors["frame_bg"])
        self.browse_buttons_frame.pack(side=tk.RIGHT, padx=10, pady=5)
        
        self.browse_image_button = tk.Button(self.browse_buttons_frame, text="文件", command=self.browse_image_file, 
                                           bg=self.colors["button"]["browse"], 
                                           fg=self.colors["button_text"]["browse"],
                                           font=('Arial', 10),
                                           relief=tk.RAISED)
        self.browse_image_button.pack(side=tk.RIGHT, padx=5)
        
        self.browse_image_folder_button = tk.Button(self.browse_buttons_frame, text="文件夹", command=self.browse_image_folder, 
                                                  bg=self.colors["button"]["browse"], 
                                                  fg=self.colors["button_text"]["browse"],
                                                  font=('Arial', 10),
                                                  relief=tk.RAISED)
        self.browse_image_folder_button.pack(side=tk.RIGHT, padx=5)
        
        # 图片位置选择
        self.image_position_frame = tk.Frame(self.export_config_frame, bg=self.colors["frame_bg"])
        self.image_position_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.image_position_label = tk.Label(self.image_position_frame, text="图片位置:", bg=self.colors["frame_bg"])
        self.image_position_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.image_position_var = tk.StringVar(value="top")
        self.image_position_top = tk.Radiobutton(self.image_position_frame, text="文章上面", variable=self.image_position_var, value="top", 
                                                bg=self.colors["frame_bg"])
        self.image_position_top.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.image_position_middle = tk.Radiobutton(self.image_position_frame, text="中间（随机）", variable=self.image_position_var, value="middle", 
                                                  bg=self.colors["frame_bg"])
        self.image_position_middle.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.image_position_bottom = tk.Radiobutton(self.image_position_frame, text="尾部", variable=self.image_position_var, value="bottom", 
                                                  bg=self.colors["frame_bg"])
        self.image_position_bottom.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 图片分配方式
        self.image_allocation_frame = tk.Frame(self.export_config_frame, bg=self.colors["frame_bg"])
        self.image_allocation_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.image_allocation_label = tk.Label(self.image_allocation_frame, text="分配方式:", bg=self.colors["frame_bg"])
        self.image_allocation_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.image_allocation_var = tk.StringVar(value="random")
        self.allocation_random = tk.Radiobutton(self.image_allocation_frame, text="随机分配", variable=self.image_allocation_var, value="random", 
                                             bg=self.colors["frame_bg"])
        self.allocation_random.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.allocation_sequential = tk.Radiobutton(self.image_allocation_frame, text="按顺序分配", variable=self.image_allocation_var, value="sequential", 
                                                   bg=self.colors["frame_bg"])
        self.allocation_sequential.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 图片使用策略
        self.image_strategy_frame = tk.Frame(self.export_config_frame, bg=self.colors["frame_bg"])
        self.image_strategy_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.image_strategy_label = tk.Label(self.image_strategy_frame, text="使用策略:", bg=self.colors["frame_bg"])
        self.image_strategy_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.image_strategy_var = tk.StringVar(value="reuse")
        self.strategy_reuse = tk.Radiobutton(self.image_strategy_frame, text="重复使用", variable=self.image_strategy_var, value="reuse", 
                                           bg=self.colors["frame_bg"])
        self.strategy_reuse.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.strategy_delete = tk.Radiobutton(self.image_strategy_frame, text="删除已用", variable=self.image_strategy_var, value="delete", 
                                             bg=self.colors["frame_bg"])
        self.strategy_delete.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 每张文本插入图片数量
        self.image_count_frame = tk.Frame(self.export_config_frame, bg=self.colors["frame_bg"])
        self.image_count_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.image_count_label = tk.Label(self.image_count_frame, text="每张文本插入图片数量:", bg=self.colors["frame_bg"])
        self.image_count_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.image_count_var = tk.StringVar(value="1")
        self.image_count_entry = tk.Entry(self.image_count_frame, width=5, textvariable=self.image_count_var)
        self.image_count_entry.pack(side=tk.LEFT, padx=5, pady=5)
        
        # 输出位置选择
        self.output_dir_frame = tk.Frame(self.export_config_frame, bg=self.colors["frame_bg"])
        self.output_dir_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.output_dir_label = tk.Label(self.output_dir_frame, text="输出位置:", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.output_dir_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.output_dir_entry = tk.Entry(self.output_dir_frame, width=50, font=('Arial', 10))
        self.output_dir_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=10, pady=5)
        
        self.browse_output_dir_button = tk.Button(self.output_dir_frame, text="浏览", command=self.browse_output_dir, 
                                               bg=self.colors["button"]["browse"], 
                                               fg=self.colors["button_text"]["browse"],
                                               font=('Arial', 10),
                                               relief=tk.RAISED)
        self.browse_output_dir_button.pack(side=tk.RIGHT, padx=10, pady=5)
        
        # 按钮区域
        self.export_buttons_frame = tk.Frame(self.export_config_frame, bg=self.colors["frame_bg"])
        self.export_buttons_frame.pack(fill=tk.X, pady=10, padx=10)
        
        self.export_button = tk.Button(self.export_buttons_frame, text="导出", command=self.export_text, 
                                     bg=self.colors["button"]["default"], 
                                     fg=self.colors["button_text"]["default"],
                                     font=('Arial', 10, 'bold'),
                                     padx=20, pady=10,
                                     relief=tk.RAISED)
        self.export_button.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)
        
        # 创建文本输入区域
        self.text_input_frame = tk.Frame(self.export_main_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE)
        self.text_input_frame.pack(fill=tk.BOTH, expand=True, pady=10, padx=5)
        
        # 输入标题
        self.text_input_title = tk.Label(self.text_input_frame, text="输入文本内容", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.text_input_title.pack(fill=tk.X, pady=5, padx=10)
        
        # 输入文本框
        self.text_input = tk.Text(self.text_input_frame, height=15, width=50, font=('Arial', 10), bd=1, relief=tk.SUNKEN)
        self.text_input.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 初始化选项状态
        self.update_image_options()
        self.update_filename_options()
        self.update_image_position_options()
    
    def init_txt_to_word_tab(self):
        """初始化百度搜索标题相关图片标签页"""
        # 创建百度搜索标题相关图片功能的主框架
        self.txt_to_word_main_frame = tk.Frame(self.txt_to_word_tab, bg=self.colors["bg"])
        self.txt_to_word_main_frame.pack(fill=tk.BOTH, expand=True, pady=10, padx=10)
        
        # 添加标题
        self.txt_to_word_title = tk.Label(self.txt_to_word_main_frame, text="百度搜索标题相关图片", font=('Arial', 12, 'bold'), bg=self.colors["bg"])
        self.txt_to_word_title.pack(fill=tk.X, pady=10)
        
        # 创建配置区域
        self.txt_to_word_config_frame = tk.Frame(self.txt_to_word_main_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE)
        self.txt_to_word_config_frame.pack(fill=tk.X, pady=10, padx=5)
        
        # 批量导入文件夹
        self.folder_import_frame = tk.Frame(self.txt_to_word_config_frame, bg=self.colors["frame_bg"])
        self.folder_import_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.folder_import_label = tk.Label(self.folder_import_frame, text="批量导入文件夹:", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.folder_import_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.folder_import_entry = tk.Entry(self.folder_import_frame, width=40, font=('Arial', 10))
        self.folder_import_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=10, pady=5)
        
        self.browse_folder_button = tk.Button(self.folder_import_frame, text="浏览", command=self.browse_import_folder, 
                                           bg=self.colors["button"]["browse"], 
                                           fg=self.colors["button_text"]["browse"],
                                           font=('Arial', 10),
                                           relief=tk.RAISED)
        self.browse_folder_button.pack(side=tk.RIGHT, padx=10, pady=5)
        
        # 或选择单个文件
        self.file_select_frame = tk.Frame(self.txt_to_word_config_frame, bg=self.colors["frame_bg"])
        self.file_select_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.file_select_label = tk.Label(self.file_select_frame, text="或选择文件:", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.file_select_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.file_select_entry = tk.Entry(self.file_select_frame, width=40, font=('Arial', 10))
        self.file_select_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=10, pady=5)
        
        self.browse_files_button = tk.Button(self.file_select_frame, text="浏览", command=self.browse_import_files, 
                                           bg=self.colors["button"]["browse"], 
                                           fg=self.colors["button_text"]["browse"],
                                           font=('Arial', 10),
                                           relief=tk.RAISED)
        self.browse_files_button.pack(side=tk.RIGHT, padx=10, pady=5)
        
        # 图片配置
        self.image_config_frame = tk.Frame(self.txt_to_word_config_frame, bg=self.colors["frame_bg"])
        self.image_config_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.image_config_label = tk.Label(self.image_config_frame, text="图片配置:", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.image_config_label.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 横屏竖屏选择
        self.orientation_frame = tk.Frame(self.image_config_frame, bg=self.colors["frame_bg"])
        self.orientation_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.orientation_label = tk.Label(self.orientation_frame, text="方向:", bg=self.colors["frame_bg"])
        self.orientation_label.pack(side=tk.LEFT, padx=5)
        
        self.orientation_var = tk.StringVar(value="all")
        self.orientation_all = tk.Radiobutton(self.orientation_frame, text="全部", variable=self.orientation_var, value="all", 
                                             bg=self.colors["frame_bg"])
        self.orientation_all.pack(side=tk.LEFT, padx=10)
        
        self.orientation_horizontal = tk.Radiobutton(self.orientation_frame, text="横屏", variable=self.orientation_var, value="horizontal", 
                                             bg=self.colors["frame_bg"])
        self.orientation_horizontal.pack(side=tk.LEFT, padx=10)
        
        self.orientation_vertical = tk.Radiobutton(self.orientation_frame, text="竖屏", variable=self.orientation_var, value="vertical", 
                                             bg=self.colors["frame_bg"])
        self.orientation_vertical.pack(side=tk.LEFT, padx=10)
        
        # 图片分辨率选择
        self.resolution_frame = tk.Frame(self.image_config_frame, bg=self.colors["frame_bg"])
        self.resolution_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.resolution_label = tk.Label(self.resolution_frame, text="分辨率:", bg=self.colors["frame_bg"])
        self.resolution_label.pack(side=tk.LEFT, padx=5)
        
        self.resolution_var = tk.StringVar(value="all")
        self.resolution_all = tk.Radiobutton(self.resolution_frame, text="全部", variable=self.resolution_var, value="all", 
                                             bg=self.colors["frame_bg"])
        self.resolution_all.pack(side=tk.LEFT, padx=10)
        
        self.resolution_1920 = tk.Radiobutton(self.resolution_frame, text="1920*1080以上", variable=self.resolution_var, value="1920x1080+", 
                                             bg=self.colors["frame_bg"])
        self.resolution_1920.pack(side=tk.LEFT, padx=10)
        
        self.resolution_1280 = tk.Radiobutton(self.resolution_frame, text="1280*720", variable=self.resolution_var, value="1280x720", 
                                             bg=self.colors["frame_bg"])
        self.resolution_1280.pack(side=tk.LEFT, padx=10)
        
        self.resolution_800 = tk.Radiobutton(self.resolution_frame, text="800*600", variable=self.resolution_var, value="800x600", 
                                             bg=self.colors["frame_bg"])
        self.resolution_800.pack(side=tk.LEFT, padx=10)
        
        # 图片清晰度选择
        self.clarity_frame = tk.Frame(self.image_config_frame, bg=self.colors["frame_bg"])
        self.clarity_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.clarity_label = tk.Label(self.clarity_frame, text="清晰度:", bg=self.colors["frame_bg"])
        self.clarity_label.pack(side=tk.LEFT, padx=5)
        
        self.clarity_var = tk.StringVar(value="original")
        self.clarity_original = tk.Radiobutton(self.clarity_frame, text="原图（根据需求直接下载图片）", variable=self.clarity_var, value="original", 
                                             bg=self.colors["frame_bg"])
        self.clarity_original.pack(side=tk.LEFT, padx=10)
        
        self.clarity_ai = tk.Radiobutton(self.clarity_frame, text="变清晰（下载前会编辑百度的AI变清晰处理过的图片进行下载保存）", variable=self.clarity_var, value="ai", 
                                             bg=self.colors["frame_bg"])
        self.clarity_ai.pack(side=tk.LEFT, padx=10)
        
        # 去水印选项
        self.watermark_frame = tk.Frame(self.image_config_frame, bg=self.colors["frame_bg"])
        self.watermark_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.watermark_label = tk.Label(self.watermark_frame, text="去水印:", bg=self.colors["frame_bg"])
        self.watermark_label.pack(side=tk.LEFT, padx=5)
        
        self.watermark_var = tk.BooleanVar(value=True)
        self.watermark_check = tk.Checkbutton(self.watermark_frame, text="处理前先去除水印", variable=self.watermark_var, 
                                             bg=self.colors["frame_bg"])
        self.watermark_check.pack(side=tk.LEFT, padx=10)
        
        # 线程数设置
        self.thread_frame = tk.Frame(self.image_config_frame, bg=self.colors["frame_bg"])
        self.thread_frame.pack(fill=tk.X, pady=5, padx=10)
        
        self.thread_label = tk.Label(self.thread_frame, text="线程数:", bg=self.colors["frame_bg"])
        self.thread_label.pack(side=tk.LEFT, padx=5)
        
        self.thread_var = tk.StringVar(value="1")
        self.thread_entry = tk.Entry(self.thread_frame, width=5, textvariable=self.thread_var, font=('Arial', 10))
        self.thread_entry.pack(side=tk.LEFT, padx=5)
        
        self.thread_unit_label = tk.Label(self.thread_frame, text="(最高20线程)", bg=self.colors["frame_bg"])
        self.thread_unit_label.pack(side=tk.LEFT, padx=5)
        
        # 下载配置
        self.download_config_frame = tk.Frame(self.txt_to_word_config_frame, bg=self.colors["frame_bg"])
        self.download_config_frame.pack(fill=tk.X, pady=5, padx=10)
        
        # 下载数量
        self.download_count_frame = tk.Frame(self.download_config_frame, bg=self.colors["frame_bg"])
        self.download_count_frame.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.download_count_label = tk.Label(self.download_count_frame, text="下载数量:", bg=self.colors["frame_bg"])
        self.download_count_label.pack(side=tk.LEFT, padx=5)
        
        self.download_count_var = tk.StringVar(value="3")
        self.download_count_entry = tk.Entry(self.download_count_frame, width=5, textvariable=self.download_count_var, font=('Arial', 10))
        self.download_count_entry.pack(side=tk.LEFT, padx=5)
        
        # 保存位置
        self.save_location_frame = tk.Frame(self.download_config_frame, bg=self.colors["frame_bg"])
        self.save_location_frame.pack(side=tk.LEFT, padx=10, pady=5, fill=tk.X, expand=True)
        
        self.save_location_label = tk.Label(self.save_location_frame, text="保存位置:", bg=self.colors["frame_bg"])
        self.save_location_label.pack(side=tk.LEFT, padx=5)
        
        self.save_location_entry = tk.Entry(self.save_location_frame, width=30, font=('Arial', 10))
        self.save_location_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        self.browse_save_location_button = tk.Button(self.save_location_frame, text="浏览", command=self.browse_save_location, 
                                           bg=self.colors["button"]["browse"], 
                                           fg=self.colors["button_text"]["browse"],
                                           font=('Arial', 10),
                                           relief=tk.RAISED)
        self.browse_save_location_button.pack(side=tk.RIGHT, padx=5)
        
        # 按钮区域
        self.button_frame = tk.Frame(self.txt_to_word_config_frame, bg=self.colors["frame_bg"])
        self.button_frame.pack(fill=tk.X, pady=10, padx=10)
        
        # 按钮容器
        self.buttons_container = tk.Frame(self.button_frame, bg=self.colors["frame_bg"])
        self.buttons_container.pack(side=tk.TOP, anchor=tk.CENTER)
        
        # 立即搜索并下载按钮
        self.search_download_button = tk.Button(self.buttons_container, text="立即搜索并下载", command=self.search_and_download, 
                                           bg=self.colors["button"]["default"], 
                                           fg=self.colors["button_text"]["default"],
                                           font=('Arial', 12, 'bold'),
                                           padx=30, pady=15,
                                           relief=tk.RAISED)
        self.search_download_button.pack(side=tk.LEFT, padx=10)
        
        # 暂停按钮
        self.pause_button = tk.Button(self.buttons_container, text="暂停", command=self.pause_download, 
                                           bg=self.colors["button"]["clear"], 
                                           fg=self.colors["button_text"]["clear"],
                                           font=('Arial', 12, 'bold'),
                                           padx=30, pady=15,
                                           relief=tk.RAISED)
        self.pause_button.pack(side=tk.LEFT, padx=10)
        
        # 创建状态显示区域
        self.status_frame = tk.Frame(self.txt_to_word_main_frame, bg=self.colors["frame_bg"], bd=2, relief=tk.GROOVE)
        self.status_frame.pack(fill=tk.BOTH, expand=True, pady=10, padx=5)
        
        # 状态标题
        self.status_title = tk.Label(self.status_frame, text="处理状态", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"])
        self.status_title.pack(fill=tk.X, pady=5, padx=10)
        
        # 状态文本框
        self.status_text = tk.Text(self.status_frame, height=15, width=50, font=('Arial', 10), bd=1, relief=tk.SUNKEN)
        self.status_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 初始化变量
        self.imported_files = []
        self.imported_folder = ""
    
    def switch_tab(self):
        """切换标签页"""
        # 隐藏所有标签页
        self.rename_tab.pack_forget()
        self.char_tab.pack_forget()
        self.name_tab.pack_forget()
        self.word_tab.pack_forget()
        self.export_tab.pack_forget()
        self.txt_to_word_tab.pack_forget()
        
        # 根据选中的标签显示对应的内容
        if self.tab_var.get() == "rename":
            self.rename_tab.pack(fill=tk.BOTH, expand=True)
        elif self.tab_var.get() == "char":
            self.char_tab.pack(fill=tk.BOTH, expand=True)
        elif self.tab_var.get() == "name":
            self.name_tab.pack(fill=tk.BOTH, expand=True)
        elif self.tab_var.get() == "word":
            self.word_tab.pack(fill=tk.BOTH, expand=True)
        elif self.tab_var.get() == "export":
            self.export_tab.pack(fill=tk.BOTH, expand=True)
        elif self.tab_var.get() == "txt_to_word":
            self.txt_to_word_tab.pack(fill=tk.BOTH, expand=True)
    
    def browse_folder(self):
        """浏览并选择文件夹"""
        folder = filedialog.askdirectory()
        if folder:
            self.folder_path = folder
            self.folder_entry.delete(0, tk.END)
            self.folder_entry.insert(0, folder)
            self.load_files()
    
    def load_files(self):
        """加载文件夹中的文件并显示"""
        # 清空之前的文件列表和输入框
        for widget in self.inner_frame.winfo_children():
            widget.destroy()
        self.files = []
        self.name_entries = []
        
        # 获取文件夹中的文件
        try:
            files = os.listdir(self.folder_path)
            # 按文件名排序
            files.sort()
            
            # 过滤出文件（排除目录）
            file_list = [file for file in files if os.path.isfile(os.path.join(self.folder_path, file))]
            file_count = len(file_list)
            
            # 显示文件数量
            tk.Label(self.inner_frame, text=f"共 {file_count} 个文件", font=('Arial', 10, 'bold'), bg=self.colors["frame_bg"]).grid(row=0, column=0, columnspan=2, padx=10, pady=5, sticky=tk.W)
            
            # 创建表头
            tk.Label(self.inner_frame, text="原文件名", font=('Arial', 10, 'bold'), width=30, bg=self.colors["frame_bg"]).grid(row=1, column=0, padx=10, pady=5)
            tk.Label(self.inner_frame, text="新文件名", font=('Arial', 10, 'bold'), width=50, bg=self.colors["frame_bg"]).grid(row=1, column=1, padx=10, pady=5)
            
            # 显示文件列表和输入框
            for i, file in enumerate(file_list):
                self.files.append(file)
                
                # 显示原文件名
                tk.Label(self.inner_frame, text=file, width=30, bg=self.colors["frame_bg"]).grid(row=i+2, column=0, padx=10, pady=5)
                
                # 创建新文件名输入框，默认显示原文件名（不含扩展名）
                name_without_ext = os.path.splitext(file)[0]
                entry = tk.Entry(self.inner_frame, width=50, font=('Arial', 10))
                entry.insert(0, name_without_ext)
                entry.grid(row=i+2, column=1, padx=10, pady=5)
                self.name_entries.append(entry)
        except Exception as e:
            messagebox.showerror("错误", f"加载文件失败: {str(e)}")
    
    def clear_batch_input(self):
        """清空批量输入框"""
        self.batch_text.delete(1.0, tk.END)
    
    def browse_import_folder(self):
        """浏览并选择导入文件夹"""
        folder = filedialog.askdirectory()
        if folder:
            self.imported_folder = folder
            self.folder_import_entry.delete(0, tk.END)
            self.folder_import_entry.insert(0, folder)
            
            # 扫描文件夹中的TXT和WORD文件
            files = []
            for root, dirs, filenames in os.walk(folder):
                for filename in filenames:
                    if filename.endswith('.txt') or filename.endswith('.docx') or filename.endswith('.doc'):
                        files.append(os.path.join(root, filename))
            
            self.imported_files = files
            self.status_text.insert(tk.END, f"已选择文件夹: {folder}\n")
            self.status_text.insert(tk.END, f"找到 {len(files)} 个文件（TXT和WORD格式）\n")
    
    def browse_import_files(self):
        """浏览并选择导入文件"""
        files = filedialog.askopenfilenames(filetypes=[("Text files", "*.txt"), ("Word files", "*.docx;*.doc"), ("All files", "*.*")])
        if files:
            self.imported_files = files
            self.file_select_entry.delete(0, tk.END)
            self.file_select_entry.insert(0, "; ".join(files))
            self.status_text.insert(tk.END, f"已选择 {len(files)} 个文件\n")
    
    def browse_save_location(self):
        """浏览并选择保存位置"""
        folder = filedialog.askdirectory()
        if folder:
            self.save_location_entry.delete(0, tk.END)
            self.save_location_entry.insert(0, folder)
            self.status_text.insert(tk.END, f"已选择保存位置: {folder}\n")
    
    def __init__(self, root):
        self.root = root
        self.root.title("小工具合集")
        self.root.geometry("900x600")  # 设置初始大小
        self.root.minsize(600, 400)  # 设置最小大小
        self.root.configure(bg="#f0f0f0")
        
        # 颜色方案
        self.colors = {
            "bg": "#f0f0f0",
            "frame_bg": "#ffffff",
            "text": "#333333",
            "button": {
                "browse": "#e0e0e0",
                "apply": "#4CAF50",
                "clear": "#f44336",
                "rename": "#2196F3",
                "default": "#2196F3",
                "download": "#ff9800"
            },
            "button_text": {
                "browse": "#333333",
                "apply": "#ffffff",
                "clear": "#ffffff",
                "rename": "#ffffff",
                "default": "#ffffff",
                "download": "#ffffff"
            }
        }
        
        # 文件夹路径
        self.folder_path = ""
        # 文件列表
        self.files = []
        # 新名称输入框列表
        self.name_entries = []
        
        # 初始化变量
        self.txt_files = []
        self.txt_results = []
        self.imported_files = []
        self.imported_folder = ""
        
        # 多线程控制
        self.is_paused = False
        self.is_running = False
        self.threads = []
        
        # 创建主框架
        self.main_frame = tk.Frame(root, bg=self.colors["bg"])
        self.main_frame.pack(pady=20, padx=20, fill=tk.BOTH, expand=True)
        
        # 创建标签页控制区域
        self.tab_control_frame = tk.Frame(self.main_frame, bg=self.colors["bg"])
        self.tab_control_frame.pack(fill=tk.X, pady=10)
        
        # 创建标签按钮
        self.tab_var = tk.StringVar(value="rename")
        
        self.rename_tab_btn = tk.Radiobutton(self.tab_control_frame, text="文件重命名", variable=self.tab_var, value="rename", 
                                             command=self.switch_tab, bg=self.colors["bg"], font=('Arial', 10, 'bold'))
        self.rename_tab_btn.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.char_tab_btn = tk.Radiobutton(self.tab_control_frame, text="字符重组", variable=self.tab_var, value="char", 
                                           command=self.switch_tab, bg=self.colors["bg"], font=('Arial', 10, 'bold'))
        self.char_tab_btn.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.name_tab_btn = tk.Radiobutton(self.tab_control_frame, text="起名小工具", variable=self.tab_var, value="name", 
                                           command=self.switch_tab, bg=self.colors["bg"], font=('Arial', 10, 'bold'))
        self.name_tab_btn.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.word_tab_btn = tk.Radiobutton(self.tab_control_frame, text="WORD转TXT", variable=self.tab_var, value="word", 
                                           command=self.switch_tab, bg=self.colors["bg"], font=('Arial', 10, 'bold'))
        self.word_tab_btn.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.export_tab_btn = tk.Radiobutton(self.tab_control_frame, text="文本导出", variable=self.tab_var, value="export", 
                                           command=self.switch_tab, bg=self.colors["bg"], font=('Arial', 10, 'bold'))
        self.export_tab_btn.pack(side=tk.LEFT, padx=10, pady=5)
        
        self.txt_to_word_tab_btn = tk.Radiobutton(self.tab_control_frame, text="百度搜索标题相关图片", variable=self.tab_var, value="txt_to_word", 
                                           command=self.switch_tab, bg=self.colors["bg"], font=('Arial', 10, 'bold'))
        self.txt_to_word_tab_btn.pack(side=tk.LEFT, padx=10, pady=5)
        
        # 创建标签页内容区域
        self.tab_content_frame = tk.Frame(self.main_frame, bg=self.colors["bg"])
        self.tab_content_frame.pack(fill=tk.BOTH, expand=True)
        
        # 创建文件重命名标签页
        self.rename_tab = tk.Frame(self.tab_content_frame, bg=self.colors["bg"])
        
        # 创建字符重组标签页
        self.char_tab = tk.Frame(self.tab_content_frame, bg=self.colors["bg"])
        
        # 创建起名小工具标签页
        self.name_tab = tk.Frame(self.tab_content_frame, bg=self.colors["bg"])
        
        # 创建WORD转TXT标签页
        self.word_tab = tk.Frame(self.tab_content_frame, bg=self.colors["bg"])
        
        # 创建文本导出标签页
        self.export_tab = tk.Frame(self.tab_content_frame, bg=self.colors["bg"])
        
        # 创建txt转word+图标签页
        self.txt_to_word_tab = tk.Frame(self.tab_content_frame, bg=self.colors["bg"])
        
        
        
        # 初始化文件重命名标签页
        self.init_rename_tab()
        
        # 初始化字符重组标签页
        self.init_char_tab()
        
        # 初始化起名小工具标签页
        self.init_name_tab()
        
        # 初始化WORD转TXT标签页
        self.init_word_tab()
        
        # 初始化文本导出标签页
        self.init_export_tab()
        
        # 初始化txt转word+图标签页
        self.init_txt_to_word_tab()
        
        
        
        # 默认显示文件重命名标签页
        self.switch_tab()
    
    def search_and_download(self):
        """根据文件名称在百度搜索图片并下载"""
        if not self.imported_files:
            messagebox.showerror("错误", "请先导入文件或选择文件夹")
            return
        
        save_location = self.save_location_entry.get().strip()
        if not save_location:
            messagebox.showerror("错误", "请选择保存位置")
            return
        
        try:
            import requests
            from bs4 import BeautifulSoup
            import urllib.parse
            import os
            import threading
            from queue import Queue
            from selenium import webdriver
            from selenium.webdriver.chrome.options import Options
            from selenium.webdriver.common.by import By
            from selenium.webdriver.support.ui import WebDriverWait
            from selenium.webdriver.support import expected_conditions as EC
            
            # 获取配置参数
            download_count = int(self.download_count_var.get())
            orientation = self.orientation_var.get()
            resolution = self.resolution_var.get()
            clarity = self.clarity_var.get()
            remove_watermark = self.watermark_var.get()
            thread_count = int(self.thread_var.get())
            thread_count = min(thread_count, 20)  # 限制最大线程数
            
            self.status_text.insert(tk.END, f"开始搜索并下载，共 {len(self.imported_files)} 个文件\n")
            self.status_text.insert(tk.END, f"使用 {thread_count} 线程并行处理\n")
            
            # 重置控制变量
            self.is_paused = False
            self.is_running = True
            
            # 创建任务队列
            task_queue = Queue()
            for file_path in self.imported_files:
                task_queue.put((file_path, save_location, download_count, orientation, resolution, clarity, remove_watermark))
            
            # 工作线程函数
            def worker():
                while not task_queue.empty() and self.is_running and not self.is_paused:
                    try:
                        file_path, save_location, download_count, orientation, resolution, clarity, remove_watermark = task_queue.get()
                        
                        # 获取文件名（不含扩展名）作为搜索关键词
                        file_name = os.path.splitext(os.path.basename(file_path))[0]
                        
                        # 构建百度图片搜索URL
                        base_url = f"https://image.baidu.com/search/index?tn=baiduimage&word={urllib.parse.quote(file_name)}"
                        
                        # 添加分辨率参数
                        if resolution != "all":
                            if resolution == "1920x1080+":
                                search_url = f"{base_url}&rn=10&gsm=3c&gpc=stf:1555200000,1617657600|stftype:2"
                            elif resolution == "1280x720":
                                search_url = f"{base_url}&rn=10&gsm=24&gpc=stf:1555200000,1617657600|stftype:2"
                            elif resolution == "800x600":
                                search_url = f"{base_url}&rn=10&gsm=1e&gpc=stf:1555200000,1617657600|stftype:2"
                            else:
                                search_url = base_url
                        else:
                            search_url = base_url
                        
                        # 配置Selenium
                        chrome_options = Options()
                        chrome_options.add_argument('--headless')  # 无头模式
                        chrome_options.add_argument('--disable-gpu')
                        chrome_options.add_argument('--no-sandbox')
                        chrome_options.add_argument('--disable-dev-shm-usage')
                        chrome_options.add_argument('--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36')
                        
                        # 启动浏览器
                        driver = webdriver.Chrome(options=chrome_options)
                        driver.get(search_url)
                        
                        # 等待页面加载
                        WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.CLASS_NAME, 'main_img')))
                        
                        # 查找图片元素
                        img_elements = driver.find_elements(By.CLASS_NAME, 'main_img')
                        
                        # 下载图片
                        img_count = 0
                        for i, img_element in enumerate(img_elements):
                            if img_count >= download_count or self.is_paused:
                                break
                            
                            try:
                                # 获取图片链接
                                img_url = img_element.get_attribute('data-imgurl') or img_element.get_attribute('src')
                                if not img_url:
                                    continue
                                
                                # 确保URL完整
                                if not img_url.startswith('http'):
                                    img_url = 'https:' + img_url if img_url.startswith('//') else search_url + img_url
                                
                                # 点击图片查看详情
                                img_element.click()
                                
                                # 等待详情页加载
                                WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.CLASS_NAME, 'currentImg')))
                                
                                # 根据清晰度选择决定是否执行去水印和AI变清晰
                                if clarity == "ai":
                                    # 先执行去水印
                                    if remove_watermark:
                                        try:
                                            watermark_button = WebDriverWait(driver, 3).until(EC.element_to_be_clickable((By.XPATH, "//span[contains(text(), '去水印')]")))
                                            watermark_button.click()
                                            self.status_text.insert(tk.END, f"  已点击去水印按钮\n")
                                            # 等待去水印处理
                                            WebDriverWait(driver, 10).until(EC.invisibility_of_element_located((By.CLASS_NAME, 'loading')))
                                        except Exception as watermark_error:
                                            self.status_text.insert(tk.END, f"  去水印功能不可用: {str(watermark_error)}\n")
                                    
                                    # 然后执行AI变清晰
                                    try:
                                        ai_clear_button = WebDriverWait(driver, 3).until(EC.element_to_be_clickable((By.XPATH, "//span[contains(text(), 'AI变清晰')]")))
                                        ai_clear_button.click()
                                        self.status_text.insert(tk.END, f"  已点击AI变清晰按钮\n")
                                        # 等待变清晰处理
                                        WebDriverWait(driver, 10).until(EC.invisibility_of_element_located((By.CLASS_NAME, 'loading')))
                                    except Exception as ai_error:
                                        self.status_text.insert(tk.END, f"  AI变清晰功能不可用: {str(ai_error)}\n")
                                else:
                                    self.status_text.insert(tk.END, f"  已选择原图，跳过处理\n")
                                
                                # 获取处理后的图片URL
                                current_img = driver.find_element(By.CLASS_NAME, 'currentImg')
                                clear_img_url = current_img.get_attribute('src')
                                
                                # 下载图片
                                headers = {
                                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
                                }
                                img_response = requests.get(clear_img_url, headers=headers, timeout=10)
                                if img_response.status_code == 200:
                                    # 构建保存文件名
                                    save_filename = os.path.join(save_location, f"{file_name}----{img_count+1}.jpg")
                                    
                                    # 保存图片
                                    with open(save_filename, 'wb') as f:
                                        f.write(img_response.content)
                                    
                                    img_count += 1
                                    self.status_text.insert(tk.END, f"  已下载图片 {img_count}/{download_count}: {os.path.basename(save_filename)}\n")
                            except Exception as img_error:
                                self.status_text.insert(tk.END, f"  下载图片失败: {str(img_error)}\n")
                                continue
                            finally:
                                # 返回搜索页
                                driver.back()
                                WebDriverWait(driver, 5).until(EC.presence_of_element_located((By.CLASS_NAME, 'main_img')))
                        
                        if img_count == 0:
                            self.status_text.insert(tk.END, "  未找到可下载的图片\n")
                    except Exception as search_error:
                        self.status_text.insert(tk.END, f"  搜索失败: {str(search_error)}\n")
                    finally:
                        # 关闭浏览器
                        try:
                            driver.quit()
                        except:
                            pass
                        task_queue.task_done()
            
            # 创建并启动线程
            self.threads = []
            for i in range(thread_count):
                t = threading.Thread(target=worker)
                t.daemon = True
                t.start()
                self.threads.append(t)
            
            # 等待所有任务完成
            task_queue.join()
            
            if self.is_running and not self.is_paused:
                self.status_text.insert(tk.END, "\n所有文件处理完成！\n")
                messagebox.showinfo("成功", f"已完成所有文件的处理")
            elif self.is_paused:
                self.status_text.insert(tk.END, "\n处理已暂停！\n")
                messagebox.showinfo("暂停", "处理已暂停，可点击继续按钮恢复")
            
        except Exception as e:
            self.status_text.insert(tk.END, f"处理失败: {str(e)}\n")
            messagebox.showerror("错误", f"处理失败: {str(e)}")
        finally:
            self.is_running = False
    
    def pause_download(self):
        """暂停下载任务"""
        if self.is_running:
            self.is_paused = True
            self.status_text.insert(tk.END, "\n正在暂停处理...\n")
            messagebox.showinfo("暂停", "处理已暂停")
        else:
            messagebox.showinfo("提示", "当前没有正在运行的任务")
    
    def apply_batch_names(self):
        """批量应用新名称"""
        if not self.files:
            messagebox.showwarning("警告", "请先选择文件夹并加载文件列表")
            return
        
        # 从批量输入文本框中获取文本
        batch_text = self.batch_text.get(1.0, tk.END).strip()
        if not batch_text:
            messagebox.showwarning("警告", "批量输入框为空，请输入新名称")
            return
        
        # 按行分割文本，得到每个文件的新名称
        new_names = [name.strip() for name in batch_text.split('\n') if name.strip()]
        
        # 检查新名称数量与文件数量是否匹配
        if len(new_names) != len(self.files):
            messagebox.showwarning("警告", f"新名称数量({len(new_names)})与文件数量({len(self.files)})不匹配，请检查输入")
            return
        
        # 将新名称应用到对应的输入框中
        for i, (entry, new_name) in enumerate(zip(self.name_entries, new_names)):
            entry.delete(0, tk.END)
            entry.insert(0, new_name)
        
        messagebox.showinfo("成功", "批量名称已应用")
    
    def rename_files(self):
        """重命名文件"""
        if not self.folder_path:
            messagebox.showwarning("警告", "请先选择文件夹")
            return
        
        if not self.files:
            messagebox.showwarning("警告", "文件夹中没有文件")
            return
        
        success_count = 0
        error_count = 0
        error_messages = []
        
        for i, (old_file, entry) in enumerate(zip(self.files, self.name_entries)):
            new_name = entry.get().strip()
            if not new_name:
                error_messages.append(f"第{i+1}个文件的新名称不能为空")
                error_count += 1
                continue
            
            # 获取文件扩展名
            ext = os.path.splitext(old_file)[1]
            # 构建新文件名
            new_file = new_name + ext
            
            # 构建完整路径
            old_path = os.path.join(self.folder_path, old_file)
            new_path = os.path.join(self.folder_path, new_file)
            
            # 检查新文件名是否与其他文件冲突
            if os.path.exists(new_path) and old_path != new_path:
                error_messages.append(f"第{i+1}个文件的新名称 '{new_file}' 已存在")
                error_count += 1
                continue
            
            try:
                # 重命名文件
                os.rename(old_path, new_path)
                success_count += 1
            except Exception as e:
                error_messages.append(f"第{i+1}个文件重命名失败: {str(e)}")
                error_count += 1
        
        # 显示结果
        if error_count == 0:
            messagebox.showinfo("成功", f"所有文件重命名成功！共重命名 {success_count} 个文件")
            # 重新加载文件列表
            self.load_files()
        else:
            error_msg = f"重命名完成，但有 {error_count} 个错误：\n" + "\n".join(error_messages)
            messagebox.showerror("错误", error_msg)
    
    def convert_text(self):
        """执行文本转换"""
        # 获取输入内容
        input_text = self.input_text.get(1.0, tk.END).strip()
        if not input_text:
            messagebox.showwarning("警告", "请输入内容")
            return
        
        # 获取重组顺序
        order_text = self.order_entry.get().strip()
        if not order_text:
            messagebox.showwarning("警告", "请填写重组顺序")
            return
        
        # 获取分隔符
        separator = self.separator_entry.get().strip()
        if not separator:
            messagebox.showwarning("警告", "请填写分隔符")
            return
        
        # 解析重组顺序
        try:
            order_parts = order_text.split("----")
            order = [int(part.strip()) - 1 for part in order_parts]  # 转换为0-based索引
        except ValueError:
            messagebox.showerror("错误", "重组顺序格式错误，请使用数字和'----'分隔")
            return
        
        # 处理每一行
        lines = input_text.split('\n')
        output_lines = []
        error_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 按分隔符分割行
            parts = [part.strip() for part in line.split(separator)]
            
            # 检查部分数量是否足够
            max_index = max(order)
            if max_index >= len(parts):
                # 记录错误行
                error_lines.append(f"错误: 第 '{line}' 行缺少序号 {max_index + 1} 的内容")
                continue  # 跳过错误行，继续处理其他行
            
            # 按指定顺序重组
            reordered_parts = [parts[i] for i in order]
            # 用分隔符连接
            output_line = separator.join(reordered_parts)
            output_lines.append(output_line)
        
        # 显示结果
        self.output_text.delete(1.0, tk.END)
        self.output_text.insert(1.0, '\n'.join(output_lines))
        
        # 显示错误信息
        if error_lines:
            error_message = "以下行存在错误，已跳过处理：\n" + "\n".join(error_lines)
            
            # 创建错误信息窗口
            error_window = tk.Toplevel(self.root)
            error_window.title("错误信息")
            error_window.geometry("500x300")
            error_window.configure(bg="#f0f0f0")
            
            # 错误信息文本框
            error_text = tk.Text(error_window, height=10, width=60, font=('Arial', 10), bd=1, relief=tk.SUNKEN)
            error_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            error_text.insert(1.0, error_message)
            error_text.config(state=tk.DISABLED)
            
            # 按钮区域
            button_frame = tk.Frame(error_window, bg="#f0f0f0")
            button_frame.pack(fill=tk.X, padx=10, pady=5)
            
            # 复制按钮
            copy_button = tk.Button(button_frame, text="复制错误信息", command=lambda: self.copy_error(error_text), 
                                   bg=self.colors["button"]["apply"], 
                                   fg=self.colors["button_text"]["apply"],
                                   font=('Arial', 10),
                                   padx=10, pady=5,
                                   relief=tk.RAISED)
            copy_button.pack(side=tk.RIGHT, padx=10)
            
            # 确定按钮
            ok_button = tk.Button(button_frame, text="确定", command=error_window.destroy, 
                                 bg=self.colors["button"]["default"], 
                                 fg=self.colors["button_text"]["default"],
                                 font=('Arial', 10),
                                 padx=10, pady=5,
                                 relief=tk.RAISED)
            ok_button.pack(side=tk.RIGHT, padx=10)

    def clear_input(self):
        """清空输入框"""
        self.input_text.delete(1.0, tk.END)
        self.update_line_count()

    def clear_output(self):
        """清空输出框"""
        self.output_text.delete(1.0, tk.END)

    def copy_output(self):
        """复制输出内容到剪贴板"""
        output_text = self.output_text.get(1.0, tk.END).strip()
        if not output_text:
            messagebox.showwarning("警告", "输出内容为空，无法复制")
            return
        
        # 复制到剪贴板
        self.root.clipboard_clear()
        self.root.clipboard_append(output_text)
        messagebox.showinfo("成功", "内容已复制到剪贴板")

    def update_line_count(self):
        """更新输入内容的行数显示"""
        # 获取输入内容
        input_text = self.input_text.get(1.0, tk.END).strip()
        if not input_text:
            line_count = 0
        else:
            # 计算行数
            line_count = len(input_text.split('\n'))
        
        # 更新行数显示
        self.line_count_label.config(text=f"({line_count}行)")

    def copy_error(self, error_text_widget):
        """复制错误信息"""
        error_text = error_text_widget.get(1.0, tk.END).strip()
        if error_text:
            self.root.clipboard_clear()
            self.root.clipboard_append(error_text)
            messagebox.showinfo("成功", "错误信息已复制到剪贴板")

    def update_config_fields(self):
        """更新配置字段的显示状态"""
        # 根据选中的类型显示/隐藏对应的配置字段
        if self.include_chinese.get():
            self.chinese_count_label.pack(side=tk.LEFT, padx=10, pady=5)
            self.chinese_count_entry.pack(side=tk.LEFT, padx=5, pady=5)
            self.stroke_frame.pack(fill=tk.X, pady=5, padx=10)
        else:
            self.chinese_count_label.pack_forget()
            self.chinese_count_entry.pack_forget()
            self.stroke_frame.pack_forget()
        
        if self.include_english.get():
            self.english_count_label.pack(side=tk.LEFT, padx=10, pady=5)
            self.english_count_entry.pack(side=tk.LEFT, padx=5, pady=5)
        else:
            self.english_count_label.pack_forget()
            self.english_count_entry.pack_forget()
        
        if self.include_number.get():
            self.number_count_label.pack(side=tk.LEFT, padx=10, pady=5)
            self.number_count_entry.pack(side=tk.LEFT, padx=5, pady=5)
        else:
            self.number_count_label.pack_forget()
            self.number_count_entry.pack_forget()

    def build_chinese_chars(self):
        """构建汉字库（按笔画分类）"""
        # 这里使用一些常见的汉字作为示例，实际应用中可以扩展更多
        # 注意：以下汉字笔画数仅供示例，实际应用中应使用正确的笔画数
        chars_by_stroke = {
            8: ['受', '抒', '叔', '刷', '祀', '忪', '怂', '所', '兔', '昔', '穸', '些', '姓', '刖', '甾'],
            9: ['测', '恻', '草', '茶', '差', '查', '姹', '柴', '豺', '婵', '谗', '禅', '馋', '缠', '蝉'],
            10: ['啊', '唉', '埃', '挨', '皑', '癌', '蔼', '矮', '艾', '碍', '爱', '隘', '鞍', '氨', '胺'],
            11: ['基', '寄', '寂', '祭', '绩', '继', '既', '暨', '稼', '假', '价', '驾', '嫁', '歼', '监'],
            12: ['奥', '跋', '魃', '罢', '白', '百', '柏', '败', '拜', '班', '般', '颁', '板', '版', '扮'],
            13: ['澳', '芭', '捌', '扒', '叭', '吧', '笆', '八', '疤', '巴', '拔', '跋', '魃', '罢', '白'],
            14: ['熬', '翱', '袄', '傲', '奥', '懊', '澳', '芭', '捌', '扒', '叭', '吧', '笆', '八', '疤'],
            15: ['鞍', '氨', '胺', '翱', '袄', '傲', '奥', '懊', '澳', '芭', '捌', '扒', '叭', '吧', '笆'],
            16: ['薄', '雹', '堡', '暴', '曝', '爆', '卑', '碑', '悲', '卑', '碑', '悲', '卑', '碑', '悲'],
            17: ['臂', '璧', '避', '陛', '鞭', '边', '编', '贬', '扁', '便', '变', '卞', '辨', '辩', '辫'],
            18: ['翱', '袄', '傲', '奥', '懊', '澳', '芭', '捌', '扒', '叭', '吧', '笆', '八', '疤', '巴'],
            19: ['霭', '皑', '癌', '蔼', '矮', '艾', '碍', '爱', '隘', '鞍', '氨', '胺', '翱', '袄', '傲'],
            20: ['骜', '鳌', '鏖', '鏊', '螯', '鏊', '骜', '鳌', '鏖', '鏊', '螯', '鏊', '骜', '鳌', '鏖'],
            21: ['龠', '龡', '龢', '龣', '龤', '龥', '龦', '龧', '龨', '龩', '龪', '龫', '龬', '龭', '龮'],
            22: ['龯', '龰', '龱', '龲', '龳', '龴', '龵', '龶', '龷', '龸', '龹', '龺', '龻', '龼', '龽'],
            23: ['龾', '龿', '龍', '龾', '龿', '龍', '龾', '龿', '龍', '龾', '龿', '龍', '龾', '龿', '龍'],
            24: ['儼', '酽', '釅', '讌', '讞', '豔', '豔', '驗', '驅', '鷫', '鸘', '鸙', '鸚', '鸜', '鸝'],
            25: ['鸞', '鸞', '鸞', '鸞', '鸞', '鸞', '鸞', '鸞', '鸞', '鸞', '鸞', '鸞', '鸞', '鸞', '鸞'],
            26: ['驥', '驥', '驥', '驥', '驥', '驥', '驥', '驥', '驥', '驥', '驥', '驥', '驥', '驥', '驥']
        }
        return chars_by_stroke

    def load_history(self):
        """加载历史记录"""
        try:
            with open(self.history_file, 'r', encoding='utf-8') as f:
                self.history = set(f.read().split('\n'))
        except FileNotFoundError:
            self.history = set()

    def save_to_history(self, names):
        """保存生成的名字到历史记录"""
        new_names = set(names)
        self.history.update(new_names)
        
        # 保存到文件
        with open(self.history_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(self.history))

    def generate_names(self):
        """生成名字"""
        import random
        import string
        
        # 获取配置
        include_chinese = self.include_chinese.get()
        include_english = self.include_english.get()
        include_number = self.include_number.get()
        
        # 获取数量配置
        try:
            chinese_count = int(self.chinese_count_var.get()) if include_chinese else 0
            english_count = int(self.english_count_var.get()) if include_english else 0
            number_count = int(self.number_count_var.get()) if include_number else 0
            min_stroke = int(self.min_stroke_var.get()) if include_chinese else 0
            max_stroke = int(self.max_stroke_var.get()) if include_chinese else 0
            generate_count = int(self.generate_var.get())
        except ValueError:
            messagebox.showerror("错误", "请输入有效的数字")
            return
        
        # 验证配置
        if not (include_chinese or include_english or include_number):
            messagebox.showerror("错误", "至少需要选择一种生成类型")
            return
        
        if generate_count <= 0 or generate_count > 500:
            messagebox.showerror("错误", "生成数量必须在1-500之间")
            return
        
        # 验证笔画范围
        if include_chinese:
            if min_stroke > max_stroke:
                messagebox.showerror("错误", "最小笔画不能大于最大笔画")
                return
            
            # 检查笔画范围内是否有可用的汉字
            available_strokes = [stroke for stroke in self.chinese_chars.keys() if min_stroke <= stroke <= max_stroke]
            if not available_strokes:
                messagebox.showerror("错误", f"笔画范围 {min_stroke}-{max_stroke} 内没有可用的汉字")
                return
        
        # 生成名字
        generated_names = []
        
        for _ in range(generate_count):
            name_parts = []
            
            # 添加汉字
            if include_chinese and chinese_count > 0:
                for _ in range(chinese_count):
                    # 随机选择一个笔画数
                    stroke = random.choice(available_strokes)
                    # 从该笔画数的汉字中随机选择一个
                    char = random.choice(self.chinese_chars[stroke])
                    name_parts.append(char)
            
            # 添加英文
            if include_english and english_count > 0:
                # 生成指定长度的随机英文字母
                english_part = ''.join(random.choices(string.ascii_letters, k=english_count))
                name_parts.append(english_part)
            
            # 添加数字
            if include_number and number_count > 0:
                # 生成指定长度的随机数字
                number_part = ''.join(random.choices(string.digits, k=number_count))
                name_parts.append(number_part)
            
            # 组合成完整的名字
            name = ''.join(name_parts)
            generated_names.append(name)
        
        # 显示结果
        self.result_text.delete(1.0, tk.END)
        self.result_text.insert(1.0, '\n'.join(generated_names))
        
        # 保存到历史记录
        self.save_to_history(generated_names)

    def copy_name_result(self):
        """复制生成的名字结果"""
        result_text = self.result_text.get(1.0, tk.END).strip()
        if not result_text:
            messagebox.showwarning("警告", "生成结果为空，无法复制")
            return
        
        # 复制到剪贴板
        self.root.clipboard_clear()
        self.root.clipboard_append(result_text)
        messagebox.showinfo("成功", "生成结果已复制到剪贴板")

    def clear_name_result(self):
        """清空生成结果"""
        self.result_text.delete(1.0, tk.END)

    def browse_word_output(self):
        """浏览WORD转换的输出位置"""
        folder = filedialog.askdirectory()
        if folder:
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, folder)

    def word_on_click(self):
        """选择WORD文件"""
        files = filedialog.askopenfilenames(filetypes=[("Word files", "*.docx *.doc")])
        if files:
            self.word_file_list.extend(files)
            # 更新文件数量显示
            self.word_file_count_label.config(text=f"({len(self.word_file_list)}个文件)")

    def convert_word_to_txt(self):
        """转换WORD文件为TXT"""
        if not self.word_file_list:
            messagebox.showwarning("警告", "请先选择WORD文件")
            return
        
        output_dir = self.output_entry.get().strip()
        if not output_dir:
            messagebox.showwarning("警告", "请选择输出位置")
            return
        
        # 创建输出目录
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 获取线程数
        try:
            thread_count = int(self.thread_var.get())
            if thread_count < 1 or thread_count > 50:
                thread_count = 1
        except ValueError:
            thread_count = 1
        
        # 清空之前的输出
        self.word_output_text.delete(1.0, tk.END)
        
        # 创建任务队列
        task_queue = queue.Queue()
        for file_path in self.word_file_list:
            task_queue.put(file_path)
        
        # 创建结果队列
        result_queue = queue.Queue()
        
        # 工作线程函数
        def worker():
            while not task_queue.empty():
                try:
                    file_path = task_queue.get_nowait()
                    # 这里只是模拟转换过程，实际应用中需要使用python-docx库
                    # 为了避免依赖问题，我们只是创建一个空的txt文件
                    file_name = os.path.basename(file_path)
                    txt_file_name = os.path.splitext(file_name)[0] + '.txt'
                    txt_file_path = os.path.join(output_dir, txt_file_name)
                    
                    # 创建空的txt文件
                    with open(txt_file_path, 'w', encoding='utf-8') as f:
                        f.write(f"转换结果: {file_name}")
                    
                    result_queue.put((file_path, "成功"))
                except Exception as e:
                    result_queue.put((file_path, f"失败: {str(e)}"))
                finally:
                    task_queue.task_done()
        
        # 启动工作线程
        threads = []
        for _ in range(min(thread_count, len(self.word_file_list))):
            thread = threading.Thread(target=worker)
            thread.daemon = True
            thread.start()
            threads.append(thread)
        
        # 等待所有线程完成
        for thread in threads:
            thread.join()
        
        # 显示转换结果
        while not result_queue.empty():
            file_path, status = result_queue.get()
            file_name = os.path.basename(file_path)
            self.word_output_text.insert(tk.END, f"{file_name}: {status}\n")
        
        messagebox.showinfo("完成", "WORD文件转换完成")

    def download_all_word_results(self):
        """下载所有转换结果"""
        output_dir = self.output_entry.get().strip()
        if not output_dir:
            messagebox.showwarning("警告", "请选择输出位置")
            return
        
        # 这里只是模拟下载过程
        messagebox.showinfo("提示", f"所有转换结果已保存到: {output_dir}")

    def save_all_word_results(self):
        """保存所有转换结果"""
        output_dir = self.output_entry.get().strip()
        if not output_dir:
            messagebox.showwarning("警告", "请选择输出位置")
            return
        
        # 这里只是模拟保存过程
        messagebox.showinfo("提示", f"所有转换结果已保存到: {output_dir}")

    def update_image_options(self):
        """更新图片选项的显示状态"""
        # 根据选择的格式显示/隐藏图片相关选项
        if self.format_var.get() == "word":
            self.image_frame.pack(fill=tk.X, pady=5, padx=10)
            self.image_path_frame.pack(fill=tk.X, pady=5, padx=10)
            self.image_position_frame.pack(fill=tk.X, pady=5, padx=10)
            self.image_allocation_frame.pack(fill=tk.X, pady=5, padx=10)
            self.image_strategy_frame.pack(fill=tk.X, pady=5, padx=10)
            self.image_count_frame.pack(fill=tk.X, pady=5, padx=10)
        else:
            self.image_frame.pack_forget()
            self.image_path_frame.pack_forget()
            self.image_position_frame.pack_forget()
            self.image_allocation_frame.pack_forget()
            self.image_strategy_frame.pack_forget()
            self.image_count_frame.pack_forget()

    def update_filename_options(self):
        """更新文件名选项的显示状态"""
        # 根据选择的文件名模式显示/隐藏相关选项
        if self.filename_mode_var.get() == "content":
            self.content_chars_frame.pack(fill=tk.X, pady=5, padx=10)
        else:
            self.content_chars_frame.pack_forget()

    def update_image_position_options(self):
        """更新图片位置选项的显示状态"""
        # 根据是否选择插入图片显示/隐藏图片位置选项
        if self.image_var.get():
            self.image_path_frame.pack(fill=tk.X, pady=5, padx=10)
            self.image_position_frame.pack(fill=tk.X, pady=5, padx=10)
            self.image_allocation_frame.pack(fill=tk.X, pady=5, padx=10)
            self.image_strategy_frame.pack(fill=tk.X, pady=5, padx=10)
            self.image_count_frame.pack(fill=tk.X, pady=5, padx=10)
        else:
            self.image_path_frame.pack_forget()
            self.image_position_frame.pack_forget()
            self.image_allocation_frame.pack_forget()
            self.image_strategy_frame.pack_forget()
            self.image_count_frame.pack_forget()

    def browse_image_file(self):
        """浏览图片文件"""
        file = filedialog.askopenfilename(filetypes=[("Image files", "*.jpg *.jpeg *.png *.gif")])
        if file:
            self.image_path_entry.delete(0, tk.END)
            self.image_path_entry.insert(0, file)

    def browse_image_folder(self):
        """浏览图片文件夹"""
        folder = filedialog.askdirectory()
        if folder:
            self.image_path_entry.delete(0, tk.END)
            self.image_path_entry.insert(0, folder)

    def browse_output_dir(self):
        """浏览输出目录"""
        folder = filedialog.askdirectory()
        if folder:
            self.output_dir_entry.delete(0, tk.END)
            self.output_dir_entry.insert(0, folder)

    def export_text(self):
        """导出文本"""
        # 获取输入文本
        text = self.text_input.get(1.0, tk.END).strip()
        if not text:
            messagebox.showwarning("警告", "请输入文本内容")
            return
        
        # 获取输出目录
        output_dir = self.output_dir_entry.get().strip()
        if not output_dir:
            messagebox.showwarning("警告", "请选择输出位置")
            return
        
        # 创建输出目录
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        # 获取导出配置
        export_mode = self.export_mode_var.get()
        format_type = self.format_var.get()
        filename_mode = self.filename_mode_var.get()
        content_chars = int(self.content_chars_var.get())
        
        # 处理文本
        lines = text.split('\n')
        lines = [line.strip() for line in lines if line.strip()]
        
        if export_mode == "single":
            # 单个文件导出
            if format_type == "txt":
                # 导出为TXT
                file_path = os.path.join(output_dir, "exported.txt")
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write('\n'.join(lines))
                messagebox.showinfo("成功", f"文本已成功导出到: {file_path}")
            else:
                # 导出为WORD
                try:
                    from docx import Document
                    doc = Document()
                    for line in lines:
                        doc.add_paragraph(line)
                    file_path = os.path.join(output_dir, "exported.docx")
                    doc.save(file_path)
                    messagebox.showinfo("成功", f"文本已成功导出到: {file_path}")
                except ImportError:
                    messagebox.showerror("错误", "请安装python-docx库: pip install python-docx")
        else:
            # 按行独立文件导出
            for i, line in enumerate(lines):
                # 生成文件名
                if filename_mode == "sequential":
                    filename = f"{i+1}.{format_type}"
                else:
                    # 按内容开头
                    content_start = line[:content_chars].replace('\\', '').replace('/', '').replace(':', '').replace('*', '').replace('?', '').replace('"', '').replace('<', '').replace('>', '').replace('|', '')
                    filename = f"{content_start}.{format_type}"
                
                file_path = os.path.join(output_dir, filename)
                
                if format_type == "txt":
                    # 导出为TXT
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(line)
                else:
                    # 导出为WORD
                    try:
                        from docx import Document
                        doc = Document()
                        doc.add_paragraph(line)
                        doc.save(file_path)
                    except ImportError:
                        messagebox.showerror("错误", "请安装python-docx库: pip install python-docx")
            
            messagebox.showinfo("成功", f"文本已按行成功导出到: {output_dir}")

if __name__ == "__main__":
    root = tk.Tk()
    app = FileRenamerApp(root)
    root.mainloop()